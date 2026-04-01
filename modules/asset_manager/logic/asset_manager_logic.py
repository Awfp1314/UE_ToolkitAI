# -*- coding: utf-8 -*-

"""
资产管理逻辑层 - 门面模式（Facade Pattern）

协调各子模块完成资产管理功能，负责：
- Qt 信号发射
- 子模块间的协调
- 向后兼容的公共 API

Task 10 重构：将 AssetManagerLogic 转换为门面模式。
"""

import re
import sys
import uuid
import shutil
import subprocess
import threading
from pathlib import Path
from typing import List, Optional, Dict, Any, Callable, Generator
from datetime import datetime

try:
    from pypinyin import lazy_pinyin, Style
except ImportError:
    lazy_pinyin = None
    Style = None

from PyQt6.QtCore import QObject, pyqtSignal

from core.logger import get_logger
from core.config.config_manager import ConfigManager
from core.exceptions import AssetError, ConfigError
from core.utils.config_utils import ConfigUtils
from .asset_model import Asset, AssetType, PackageType
from .asset_core import AssetCore
from .asset_scanner import AssetScanner
from .thumbnail_manager import ThumbnailManager
from .asset_preview_coordinator import AssetPreviewCoordinator
from .asset_local_config_manager import AssetLocalConfigManager
from .file_operations import FileOperations
from .search_engine import SearchEngine
from .screenshot_processor import ScreenshotProcessor
from .asset_migrator import AssetMigrator
from .config_handler import ConfigHandler
from .preview_manager import PreviewManager

logger = get_logger(__name__)


class AssetManagerLogic(QObject):
    """资产管理逻辑类 - 门面模式

    协调各子模块完成资产管理功能。所有业务逻辑委托给专门的子模块，
    本类仅负责协调调用和 Qt 信号发射。

    Signals:
        asset_added: 资产添加完成信号 (Asset)
        asset_removed: 资产删除完成信号 (str: asset_id)
        assets_loaded: 资产列表加载完成信号 (List[Asset])
        preview_started: 预览启动信号 (str: asset_id)
        preview_finished: 预览完成信号
        thumbnail_updated: 缩略图更新信号 (str: asset_id, str: thumbnail_path)
        error_occurred: 错误发生信号 (str: error_message)
        progress_updated: 进度更新信号 (int: current, int: total, str: message)
        asset_selected: 资产选中信号 (dict)
    """

    asset_added = pyqtSignal(object)
    asset_removed = pyqtSignal(str)
    assets_loaded = pyqtSignal(list)
    preview_started = pyqtSignal(str)
    preview_finished = pyqtSignal()
    thumbnail_updated = pyqtSignal(str, str)
    error_occurred = pyqtSignal(str)
    progress_updated = pyqtSignal(int, int, str)
    asset_selected = pyqtSignal(dict)
    new_asset_detected = pyqtSignal(str, object)  # path, AssetType

    def __init__(self, config_dir: Path):
        super().__init__()
        self.config_dir = Path(config_dir)

        # 获取模块自身的配置模板路径
        module_dir = Path(__file__).parent.parent
        template_path = module_dir / "config_template.json"

        # 导入配置 schema
        from ..config_schema import get_asset_manager_schema
        
        self.config_manager = ConfigManager("asset_manager",
                                           template_path=template_path,
                                           config_schema=get_asset_manager_schema())

        # ─── 子模块初始化 ──────────────────────────────────
        self._file_ops = FileOperations(logger)
        self._search_engine = SearchEngine(logger)
        self._screenshot_processor = ScreenshotProcessor(logger)
        self._asset_migrator = AssetMigrator(self._file_ops, logger)
        self._config_handler = ConfigHandler(self.config_manager, logger)
        self._preview_manager = PreviewManager(self._file_ops, logger)
        self._asset_core = AssetCore(logger)
        self._thumbnail_manager = ThumbnailManager(logger)
        self._asset_scanner = AssetScanner(logger, thumbnail_manager=self._thumbnail_manager)
        self._preview_coordinator = AssetPreviewCoordinator(
            config_manager=self.config_manager,
            file_ops=self._file_ops,
            screenshot_processor=self._screenshot_processor,
            logger=logger
        )
        self._local_config = AssetLocalConfigManager(self.config_manager, logger)

        # ─── 向后兼容的属性引用 ────────────────────────────
        self.assets = self._asset_core.assets
        self.categories = self._asset_core.categories

        self._load_config()

        logger.info("资产管理逻辑初始化完成")

    # ─── 向后兼容的属性 ────────────────────────────────────

    @property
    def local_config_path(self):
        return self._local_config.local_config_path

    @local_config_path.setter
    def local_config_path(self, value):
        self._local_config.local_config_path = value

    @property
    def thumbnails_dir(self):
        return self._local_config.thumbnails_dir

    @thumbnails_dir.setter
    def thumbnails_dir(self, value):
        self._local_config.thumbnails_dir = value

    @property
    def documents_dir(self):
        return self._local_config.documents_dir

    @documents_dir.setter
    def documents_dir(self, value):
        self._local_config.documents_dir = value

    @property
    def current_preview_process(self):
        return self._preview_coordinator.current_preview_process

    @current_preview_process.setter
    def current_preview_process(self, value):
        self._preview_coordinator.current_preview_process = value

    @property
    def current_preview_project_path(self):
        return self._preview_coordinator.current_preview_project_path

    @current_preview_project_path.setter
    def current_preview_project_path(self, value):
        self._preview_coordinator.current_preview_project_path = value


    # ─── 配置加载与保存（协调层）─────────────────────────────

    def _load_config(self) -> None:
        """加载配置 - 优先从缓存加载，异步扫描检查变化"""
        config = self._local_config.load_global_config()
        config = self._local_config.migrate_global_config(config)

        # 优先读取新格式的 current_asset_library，兼容旧格式
        asset_library_path = (config.get("current_asset_library", "")
                               or config.get("asset_library_path", ""))
        
        # 兼容旧格式的 asset_library_configs，取第一个配置的资产库路径
        if not asset_library_path:
            asset_library_configs = config.get("asset_library_configs", {})
            if asset_library_configs:
                asset_library_path = list(asset_library_configs.keys())[0]
        
        if not asset_library_path or not Path(asset_library_path).exists():
            logger.warning("资产库路径未设置或不存在，不加载任何资产")
            self.assets.clear()
            self.assets_loaded.emit(self.assets)
            return

        # 设置本地路径
        self._local_config.setup_local_paths(asset_library_path)

        # 加载本地配置
        local_config = self._local_config.load_local_config()
        if local_config:
            lib_config = local_config
        else:
            asset_library_configs = config.get("asset_library_configs", {})
            lib_config = asset_library_configs.get(asset_library_path, {})
            if lib_config:
                self._local_config.save_local_config(lib_config)

        # 加载分类（就地修改，保持与 _asset_core.categories 的引用一致）
        new_categories = lib_config.get("categories", config.get("categories", ["默认分类"]))
        print("=" * 80)
        print(f"[加载分类] lib_config.get('categories'): {lib_config.get('categories')}")
        print(f"[加载分类] config.get('categories'): {config.get('categories')}")
        print(f"[加载分类] new_categories: {new_categories}")
        print("=" * 80)
        if "默认分类" not in new_categories:
            new_categories.insert(0, "默认分类")
        self.categories.clear()
        self.categories.extend(new_categories)
        print(f"[加载分类] self.categories 更新后: {self.categories}")
        print("=" * 80)

        # 优先从缓存加载资产（快速启动）
        cached_assets_data = lib_config.get("assets", config.get("assets", []))
        
        # 调试：检查初始缓存数据
        if cached_assets_data and len(cached_assets_data) > 1:
            print(f"[DEBUG] 初始cached_assets_data[0]: {cached_assets_data[0].get('name')}, engine_min_version='{cached_assets_data[0].get('engine_min_version', 'KEY_NOT_FOUND')}'")
            print(f"[DEBUG] 初始cached_assets_data[1]: {cached_assets_data[1].get('name')}, engine_min_version='{cached_assets_data[1].get('engine_min_version', 'KEY_NOT_FOUND')}'")
        
        if cached_assets_data:
            # 修复盘符（如果资产库从 F: 移到 E: 等情况）
            current_drive = str(Path(asset_library_path).drive)
            cached_assets_data = self._asset_scanner._fix_drive_letter_in_cache(
                cached_assets_data, current_drive
            )
            
            # 调试：检查修复盘符后
            if len(cached_assets_data) > 1:
                print(f"[DEBUG] 修复盘符后[0]: {cached_assets_data[0].get('name')}, engine_min_version='{cached_assets_data[0].get('engine_min_version', 'KEY_NOT_FOUND')}'")
                print(f"[DEBUG] 修复盘符后[1]: {cached_assets_data[1].get('name')}, engine_min_version='{cached_assets_data[1].get('engine_min_version', 'KEY_NOT_FOUND')}')")
            
            # 修复资产库路径变更（如果资产库文件夹改名）
            if cached_assets_data and cached_assets_data[0].get("path"):
                # 从缓存中的第一个资产路径推断旧的资产库路径
                first_asset_path = Path(cached_assets_data[0]["path"])
                # 找到资产库根目录（假设资产在分类文件夹下）
                # 例如：D:\OldName\默认分类\Asset -> D:\OldName
                old_library_path = first_asset_path.parent.parent
                new_library_path = Path(asset_library_path)
                
                if old_library_path != new_library_path:
                    cached_assets_data = self._asset_scanner._fix_library_path_in_cache(
                        cached_assets_data, str(old_library_path), str(new_library_path)
                    )
            
            logger.info(f"从缓存快速加载 {len(cached_assets_data)} 个资产")
            self._load_assets_from_config(cached_assets_data)
            self._search_engine.build_pinyin_cache(self.assets)
        else:
            # 无缓存，同步扫描（首次使用）
            self._scan_asset_library(Path(asset_library_path), cached_assets_data)

    def _save_config(self) -> None:
        """保存配置 - 委托给 _local_config"""
        current_lib_path = self.get_asset_library_path()
        self._local_config.save_full_config(self.assets, self.categories, current_lib_path)

    def _scan_asset_library(self, library_path: Path,
                             cached_assets_data: List[Dict[str, Any]]) -> None:
        """扫描资产库 - 委托给 _asset_scanner"""
        self.assets.clear()

        scanned_assets = self._asset_scanner.scan_asset_library(
            library_path=library_path,
            cached_assets_data=cached_assets_data,
            categories=self.categories,
            thumbnails_dir=self.thumbnails_dir,
            progress_callback=lambda c, t, m: self.progress_updated.emit(c, t, m)
        )
        self.assets.extend(scanned_assets)

        logger.info(f"资产库扫描完成，共加载 {len(self.assets)} 个资产")

        self._thumbnail_manager.migrate_thumbnails_and_docs()

        # 构建搜索引擎的拼音缓存
        self._search_engine.build_pinyin_cache(self.assets)

        self.assets_loaded.emit(self.assets)

        self._save_config()
        logger.info("资产配置已保存（包含缩略图路径修复）")

    def rescan_in_background(self) -> None:
        """后台增量扫描：检测缓存外新增/删除的资产，完成后自动更新并保存"""
        from core.services import thread_service

        asset_library_path = self.get_asset_library_path()
        if not asset_library_path or not asset_library_path.exists():
            return

        cached_paths = {str(a.path) for a in self.assets}

        # 将当前内存中的资产序列化为缓存数据，供扫描器匹配使用
        # 这样已有资产会保留原始 ID、名称、缩略图等元数据
        current_assets_data = []
        for a in self.assets:
            current_assets_data.append({
                "id": a.id,
                "name": a.name,
                "asset_type": a.asset_type.value,
                "path": str(a.path),
                "category": a.category,
                "file_extension": a.file_extension,
                "thumbnail_path": str(a.thumbnail_path) if a.thumbnail_path else None,
                "thumbnail_source": a.thumbnail_source,
                "size": a.size,
                "created_time": a.created_time.isoformat() if a.created_time else None,
                "description": a.description,
                "engine_min_version": getattr(a, 'engine_min_version', ''),
            })

        def scan_task(cancel_token):
            # 扫描文件系统获取最新资产（传入当前缓存以保留元数据）
            scanned = self._asset_scanner.scan_asset_library(
                library_path=asset_library_path,
                cached_assets_data=current_assets_data,
                categories=self.categories,
                thumbnails_dir=self.thumbnails_dir,
                progress_callback=None,
            )
            scanned_paths = {str(a.path) for a in scanned}

            added = [a for a in scanned if str(a.path) not in cached_paths]
            removed_paths = cached_paths - scanned_paths

            return added, removed_paths

        def on_result(result):
            if result is None:
                return
            added, removed_paths = result
            changed = False

            if removed_paths:
                self.assets[:] = [a for a in self.assets if str(a.path) not in removed_paths]
                logger.info(f"增量扫描：移除 {len(removed_paths)} 个已删除资产")
                changed = True

            if added:
                self.assets.extend(added)
                logger.info(f"增量扫描：发现 {len(added)} 个新资产")
                changed = True

            if changed:
                self._search_engine.build_pinyin_cache(self.assets)
                self._save_config()
                self.assets_loaded.emit(self.assets)
                logger.info("增量扫描完成，已更新缓存")
            else:
                logger.info("增量扫描完成，无变化")

        thread_service.run_async(
            scan_task,
            on_result=on_result,
            on_error=lambda e: logger.warning(f"后台增量扫描失败: {e}"),
        )

    # ─── 向后兼容的委托方法（扫描/缩略图）──────────────────

    def _scan_category_folder(self, category_folder: Path, category: str,
                               cached_assets_dict: Dict[str, Dict[str, Any]]) -> List[Asset]:
        """委托给 AssetScanner"""
        return self._asset_scanner.scan_category_folder(
            category_folder, category, cached_assets_dict, self.thumbnails_dir
        )

    def _fix_drive_letter_in_cache(self, cached_assets_data: List[Dict[str, Any]],
                                    current_drive: str) -> List[Dict[str, Any]]:
        """委托给 AssetScanner"""
        return self._asset_scanner._fix_drive_letter_in_cache(cached_assets_data, current_drive)

    def _get_size(self, path: Path) -> int:
        """委托给 AssetScanner"""
        return self._asset_scanner._get_size(path)

    def _find_thumbnail_by_asset_id(self, asset_id: str) -> Optional[Path]:
        """委托给 ThumbnailManager"""
        return self._thumbnail_manager.find_thumbnail_by_asset_id(asset_id, self.thumbnails_dir)

    def _find_existing_thumbnail_for_new_asset(self, asset_path: Path) -> Optional[Path]:
        """委托给 ThumbnailManager"""
        return self._thumbnail_manager.find_existing_thumbnail_for_new_asset(
            asset_path, self.thumbnails_dir
        )

    def _restore_thumbnail_from_asset(self, asset_path: Path, asset_id: str) -> Optional[Path]:
        """委托给 ThumbnailManager"""
        return self._thumbnail_manager.restore_thumbnail_from_asset(
            asset_path, asset_id, self.thumbnails_dir
        )

    def _migrate_thumbnails_and_docs(self) -> None:
        """委托给 ThumbnailManager"""
        self._thumbnail_manager.migrate_thumbnails_and_docs()

    # ─── 向后兼容的配置委托方法 ────────────────────────────

    def _load_local_config(self) -> Optional[Dict[str, Any]]:
        """委托给 AssetLocalConfigManager"""
        return self._local_config.load_local_config()

    def _save_local_config(self, config: Dict[str, Any],
                            create_backup: bool = False) -> bool:
        """委托给 AssetLocalConfigManager"""
        return self._local_config.save_local_config(config, create_backup)

    def _migrate_config(self, config: Dict[str, Any]) -> Dict[str, Any]:
        """委托给 AssetLocalConfigManager"""
        return self._local_config.migrate_global_config(config)

    def _validate_config_before_save(self, config: Dict[str, Any]) -> bool:
        """委托给 AssetLocalConfigManager"""
        return self._local_config.validate_config_before_save(config, len(self.assets))

    def _should_create_backup(self) -> bool:
        """委托给 AssetLocalConfigManager"""
        return self._local_config.should_create_backup(len(self.assets))

    def _backup_current_config(self) -> bool:
        """委托给 AssetLocalConfigManager"""
        return self._local_config._backup_current_config()

    def _cleanup_old_backups(self, backup_dir: Path, keep_count: int = 20) -> None:
        """委托给 AssetLocalConfigManager"""
        self._local_config._cleanup_old_backups(backup_dir, keep_count)


    # ─── 资产 CRUD 操作 ────────────────────────────────────

    def add_asset(self, asset_path: Path, asset_type: AssetType, name: str = "",
                  category: str = "默认分类", description: str = "",
                  create_markdown: bool = False, engine_version: str = "",
                  package_type: PackageType = PackageType.CONTENT) -> Optional[Asset]:
        """添加资产（将资产移动到资产库）"""
        return self._add_asset_impl(
            asset_path, asset_type, name, category, description,
            create_markdown, engine_version, package_type=package_type,
            progress_callback=None
        )

    def add_asset_stream(self, **kwargs) -> Generator[str, None, None]:
        """SSE 流式资产添加接口
        
        yields:
            "data: {"progress": 10, "message": "..."}\n\n"
        """
        def stream_callback(current, total, message):
            import json
            progress = int((current / total * 100)) if total > 0 else 0
            data = json.dumps({"progress": progress, "message": message}, ensure_ascii=False)
            # 注意：这里我们无法直接从回调返回，需要某种方式传递出去
            # 在 SSE 场景下，我们通常会手动控制逻辑循环
            pass

    def add_asset_async(self, asset_path: Path, asset_type: AssetType, name: str = "",
                        category: str = "默认分类", description: str = "",
                        create_markdown: bool = False, engine_version: str = "",
                        package_type: PackageType = PackageType.CONTENT,
                        plugin_folder_name: str = "",
                        original_filename: str = "",
                        progress_callback: Optional[Callable[[int, int, str], None]] = None
                        ) -> Optional[Asset]:
        """异步添加资产（支持进度回调）
        
        Args:
            original_filename: 原始文件名（用于压缩包，避免使用临时目录名）
        """
        return self._add_asset_impl(
            asset_path, asset_type, name, category, description,
            create_markdown, engine_version, package_type=package_type,
            plugin_folder_name=plugin_folder_name,
            original_filename=original_filename,
            progress_callback=progress_callback
        )

    def _add_asset_impl(self, asset_path: Path, asset_type: AssetType, name: str,
                         category: str, description: str, create_markdown: bool,
                         engine_version: str,
                         package_type: PackageType = PackageType.CONTENT,
                         plugin_folder_name: str = "",
                         original_filename: str = "",
                         progress_callback: Optional[Callable] = None) -> Optional[Asset]:
        """添加资产的统一实现

        根据 package_type 创建不同的包装结构：
        - CONTENT:  名称/Content/...
        - PROJECT:  名称/Project/...
        - PLUGIN:   名称/Plugins/插件原名/...
        - OTHERS:   名称/Others/...

        Args:
            plugin_folder_name: 插件类型时，插件的原始文件夹名称
            original_filename: 原始文件名（用于压缩包，避免使用临时目录名）
        """
        # 性能监控：记录开始时间和各步骤耗时
        import time
        perf_start = time.time()
        perf_steps = {}

        def record_step(step_name: str, step_start_time: float):
            """记录步骤耗时"""
            elapsed = time.time() - step_start_time
            perf_steps[step_name] = elapsed
            logger.info(f"⏱️ [{step_name}] 耗时: {elapsed:.3f}秒")

        # 暂时禁用自动检测器，避免在添加资产期间触发重复检测
        auto_detector_was_enabled = False
        if hasattr(self, '_auto_detector') and self._auto_detector:
            auto_detector_was_enabled = self._auto_detector._enabled
            if auto_detector_was_enabled:
                logger.info("暂时禁用自动检测器（添加资产期间）")
                self._auto_detector.stop()

        try:
            if progress_callback:
                progress_callback(0, 100, "准备添加资产...")

            # 步骤 1: 路径验证
            step_start = time.time()
            if not asset_path.exists():
                error_msg = f"资产路径不存在: {asset_path}"
                logger.error(error_msg)
                self.error_occurred.emit(error_msg)
                return None

            library_path = self.get_asset_library_path()
            if not library_path or not library_path.exists():
                error_msg = "资产库路径未设置或不存在，请先设置资产库路径"
                logger.error(error_msg)
                self.error_occurred.emit(error_msg)
                return None
            record_step('01-路径验证', step_start)

            if progress_callback:
                progress_callback(5, 100, "创建资产目录...")

            # 步骤 2: 创建目录结构
            step_start = time.time()

            # 确保分类文件夹存在
            category_folder = library_path / category
            if not category_folder.exists():
                category_folder.mkdir(parents=True, exist_ok=True)

            # 确定资产包装文件夹名称
            asset_wrapper_name = name if name else asset_path.name

            # CONTENT 类型兜底
            if package_type == PackageType.CONTENT and asset_wrapper_name.lower() == "content":
                inferred_name = ""
                try:
                    if asset_path.is_dir() and asset_path.name.lower() == "content":
                        sub_dirs = [
                            d for d in asset_path.iterdir()
                            if d.is_dir() and not d.name.startswith('.') and d.name.lower() != "content"
                        ]
                        if sub_dirs:
                            inferred_name = sub_dirs[0].name
                    if not inferred_name and asset_path.name.lower() != "content":
                        inferred_name = asset_path.name
                except Exception:
                    inferred_name = ""

                if inferred_name:
                    logger.info(f"CONTENT 名称兜底: {asset_wrapper_name} -> {inferred_name}")
                    asset_wrapper_name = inferred_name

            wrapper_path = category_folder / asset_wrapper_name
            if wrapper_path.exists():
                counter = 1
                while wrapper_path.exists():
                    wrapper_path = category_folder / f"{asset_wrapper_name}_{counter}"
                    counter += 1
                asset_wrapper_name = wrapper_path.name

            # 创建包装结构
            wrapper_path.mkdir(parents=True, exist_ok=True)
            wrapper_folder_name = package_type.wrapper_folder
            inner_folder = wrapper_path / wrapper_folder_name
            inner_folder.mkdir(parents=True, exist_ok=True)

            record_step('02-创建目录', step_start)

            if progress_callback:
                progress_callback(10, 100, "移动资产文件...")

            # 步骤 3: 文件复制/移动
            step_start = time.time()

            def move_progress(current, total, message):
                if progress_callback and total > 0:
                    move_pct = (current / total) * 70
                    progress_callback(10 + int(move_pct), 100, f"移动文件: {message}")
                else:
                    self.progress_updated.emit(current, total, message)

            project_file_rel = ""

            if package_type == PackageType.CONTENT:
                content_source = asset_path
                is_source_content = False

                def _single_content_child(path):
                    try:
                        content_dirs = [
                            d for d in path.iterdir()
                            if d.is_dir() and d.name.lower() == "content"
                        ]
                        return content_dirs[0] if len(content_dirs) == 1 else None
                    except Exception:
                        return None

                if asset_path.is_dir():
                    current = asset_path
                    depth = 0
                    while depth < 5 and current.name.lower() == "content":
                        inner = _single_content_child(current)
                        if not inner:
                            break
                        logger.info(f"检测到嵌套 Content 包装，下钻: {current} -> {inner}")
                        current = inner
                        depth += 1

                    if current == asset_path:
                        inner = _single_content_child(current)
                        if inner:
                            logger.info(f"检测到外层目录仅包含 Content，自动定位到: {inner}")
                            current = inner

                    content_source = current
                    is_source_content = (content_source.name.lower() == "content")

                success = self._move_contents_to_folder(
                    content_source, inner_folder, is_flatten=is_source_content,
                    progress_callback=progress_callback
                )
            elif package_type == PackageType.PROJECT:
                target_path = inner_folder / asset_path.name
                logger.info(f"导入 UE 项目: {asset_path} -> {target_path}")
                success = self._file_ops.safe_copytree(
                    asset_path, target_path, progress_callback=move_progress
                )
                if success is not False:
                    for uproject in target_path.rglob("*.uproject"):
                        try:
                            project_file_rel = str(uproject.relative_to(wrapper_path))
                            logger.info(f"找到 .uproject: {project_file_rel}")
                        except ValueError:
                            project_file_rel = str(uproject)
                        break
            elif package_type == PackageType.PLUGIN:
                actual_plugin_name = plugin_folder_name if plugin_folder_name else asset_path.name
                target_path = inner_folder / actual_plugin_name
                logger.info(f"导入 UE 插件: {asset_path} -> {target_path} (插件名: {actual_plugin_name})")
                success = self._file_ops.safe_copytree(
                    asset_path, target_path, progress_callback=move_progress
                )
            else:
                # Others 类型：智能提取并分类整理
                logger.info(f"导入其他资源（智能整理模式）: {asset_path}")
                success = self._process_others_asset(
                    asset_path, inner_folder, original_filename, 
                    progress_callback, perf_steps
                )

            if success is False:
                error_msg = f"复制资产失败: {asset_path}"
                logger.error(error_msg)
                self.error_occurred.emit(error_msg)
                return None

            record_step('03-复制文件', step_start)
            logger.info(f"资产已复制到: {inner_folder} (类型: {package_type.display_name})")

            if progress_callback:
                progress_callback(80, 100, "创建资产记录...")

            # 步骤 4: 创建资产记录
            step_start = time.time()

            asset_id = str(uuid.uuid4())
            asset_name = asset_wrapper_name
            size = self._file_ops.calculate_size(wrapper_path)

            asset = Asset(
                id=asset_id,
                name=asset_name,
                asset_type=AssetType.PACKAGE,
                path=wrapper_path,
                category=category,
                package_type=package_type,
                size=size,
                description=description,
                engine_min_version=engine_version,
                project_file=project_file_rel
            )

            self.assets.append(asset)
            self._search_engine.build_pinyin_cache([asset])

            if package_type == PackageType.PLUGIN:
                self._set_plugin_default_thumbnail(asset)
            elif package_type == PackageType.OTHERS:
                self._set_others_default_thumbnail(asset, wrapper_path)

            record_step('04-创建记录', step_start)

            if progress_callback:
                progress_callback(90, 100, "保存配置...")
            else:
                self.progress_updated.emit(0, 1, "正在保存配置...")

            # 步骤 5: 保存配置
            step_start = time.time()
            self._save_config()
            record_step('05-保存配置', step_start)
            logger.info(f"添加资产成功: {asset_name} ({asset_type.value})")

            self.asset_added.emit(asset)

            if create_markdown:
                if progress_callback:
                    progress_callback(95, 100, "创建文档...")
                else:
                    self.progress_updated.emit(0, 1, "正在创建文档...")

                # 步骤 6: 创建文档
                step_start = time.time()
                self._create_asset_markdown(asset)
                record_step('06-创建文档', step_start)

                if not progress_callback:
                    self.progress_updated.emit(1, 1, "文档创建完成")

            if progress_callback:
                progress_callback(100, 100, "资产添加完成！")
            else:
                self.progress_updated.emit(1, 1, "资产添加完成！")

            # 性能监控：记录总耗时并输出到日志文件
            total_time = time.time() - perf_start
            self._log_performance(asset_name, package_type, size, perf_steps, total_time, wrapper_path)

            return asset

        except Exception as e:
            error_msg = f"添加资产失败: {e}"
            logger.error(error_msg, exc_info=True)
            self.error_occurred.emit(error_msg)
            return None
        finally:
            # 恢复自动检测器
            if auto_detector_was_enabled and hasattr(self, '_auto_detector') and self._auto_detector:
                logger.info("恢复自动检测器")
                self._auto_detector.start()


    def _move_contents_to_folder(self, source_path: Path, target_folder: Path,
                                   is_flatten: bool = False,
                                   progress_callback: Optional[Callable] = None) -> bool:
        """将源路径内容移动到目标文件夹
        
        Args:
            source_path: 源路径
            target_folder: 目标文件夹
            is_flatten: 如果为 True，将源路径的子项直接移到目标文件夹（避免嵌套）
            progress_callback: 进度回调
            
        Returns:
            bool: 成功返回 True
        """
        if is_flatten:
            # 源路径就是 Content 文件夹，逐个复制其子项到目标 Content
            items = list(source_path.iterdir())
            total_items = len(items)
            logger.info(f"源路径为 Content 文件夹，复制 {total_items} 个子项到: {target_folder}")
            
            # 预先扫描所有文件，计算总大小（用于进度映射）
            total_bytes = 0
            item_sizes = []
            for item in items:
                if item.is_dir():
                    # 递归计算目录大小
                    dir_size = sum(f.stat().st_size for f in item.rglob('*') if f.is_file())
                    item_sizes.append(dir_size)
                    total_bytes += dir_size
                else:
                    file_size = item.stat().st_size
                    item_sizes.append(file_size)
                    total_bytes += file_size
            
            logger.info(f"📊 统计到 {total_items} 个子项，总大小 {self._file_ops.format_size(total_bytes)}")
            
            # 复制每个子项
            copied_bytes = 0
            for idx, (item, item_size) in enumerate(zip(items, item_sizes), 1):
                target_item = target_folder / item.name
                
                # 定义子项的进度回调（映射到 10-80% 范围）
                def item_progress(current_bytes, total_bytes_item, message):
                    if progress_callback and total_bytes > 0:
                        # 计算当前总进度（已复制 + 当前文件进度）
                        current_total = copied_bytes + current_bytes
                        # 映射到 10-80% 范围
                        pct = (current_total / total_bytes) * 70
                        progress_callback(10 + int(pct), 100, f"复制: {item.name}")
                
                # 复制子项
                if item.is_dir():
                    item_success = self._file_ops.safe_copytree(
                        item, target_item, progress_callback=item_progress
                    )
                else:
                    item_success = self._file_ops.safe_copy_file(
                        item, target_item, progress_callback=item_progress
                    )
                
                if item_success is False:
                    logger.error(f"复制子项失败: {item} -> {target_item}")
                    return False
                
                # 更新已复制字节数
                copied_bytes += item_size
                
                # 报告子项完成进度
                if progress_callback and total_bytes > 0:
                    pct = (copied_bytes / total_bytes) * 70
                    progress_callback(10 + int(pct), 100, f"已完成: {item.name}")
            
            logger.info(f"✅ 成功复制 {total_items} 个子项，总大小 {self._file_ops.format_size(copied_bytes)}")
            return True
        else:
            # 标准流程：复制整个路径到目标文件夹内
            target_path = target_folder / source_path.name
            logger.info(f"开始复制资产: {source_path} -> {target_path}")

            def copy_progress(current, total, message):
                if progress_callback and total > 0:
                    copy_pct = (current / total) * 70
                    progress_callback(10 + int(copy_pct), 100, f"复制文件: {message}")
                else:
                    self.progress_updated.emit(current, total, message)

            if source_path.is_dir():
                return self._file_ops.safe_copytree(
                    source_path, target_path, progress_callback=copy_progress
                )
            else:
                return self._file_ops.safe_copy_file(
                    source_path, target_path, progress_callback=copy_progress
                )

    def _create_asset_markdown(self, asset: Asset) -> None:
        """创建资产的 Word 文档并打开"""
        try:
            if not self.documents_dir:
                logger.error("本地文档目录未设置")
                return

            documents_dir = self.documents_dir
            documents_dir.mkdir(parents=True, exist_ok=True)

            doc_path = documents_dir / f"{asset.id}.docx"

            # 使用 python-docx 创建 Word 文档
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            doc = Document()

            # 标题：资产信息表（12pt）
            title = doc.add_heading('资产信息表', level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title.runs[0]
            title_run.font.size = Pt(12)

            # 添加分隔线（8pt）
            sep_para = doc.add_paragraph('=' * 50)
            sep_para.runs[0].font.size = Pt(8)

            # 基本信息区（8pt）
            doc.add_paragraph()
            info_lines = [
                f"资产名称: {asset.name}",
                f"资产ID: {asset.id}",
                f"资产类型: {asset.package_type.display_name if hasattr(asset, 'package_type') else asset.asset_type.value}",
                f"分类: {asset.category}",
                f"文件路径: {asset.path}",
                f"文件大小: {self._file_ops.format_size(asset.size)}",
                f"创建时间: {asset.created_time.strftime('%Y-%m-%d %H:%M:%S') if asset.created_time else '未知'}",
            ]
            
            # 如果有引擎版本信息，添加到列表
            if hasattr(asset, 'engine_min_version') and asset.engine_min_version:
                info_lines.append(f"引擎版本: {asset.engine_min_version}")
            
            for line in info_lines:
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.size = Pt(8)

            # 分隔线（8pt）
            doc.add_paragraph()
            sep_para2 = doc.add_paragraph('=' * 50)
            sep_para2.runs[0].font.size = Pt(8)
            doc.add_paragraph()

            # 描述区（10pt 标题，8pt 内容）
            desc_heading = doc.add_heading('资产描述', level=2)
            desc_heading.runs[0].font.size = Pt(10)
            desc_content = doc.add_paragraph(asset.description or '暂无描述')
            for run in desc_content.runs:
                run.font.size = Pt(8)
            doc.add_paragraph()

            # 分隔线（8pt）
            sep_para3 = doc.add_paragraph('─' * 50)
            sep_para3.runs[0].font.size = Pt(8)
            doc.add_paragraph()

            # 使用说明区（10pt 标题，8pt 内容）
            usage_heading = doc.add_heading('使用说明', level=2)
            usage_heading.runs[0].font.size = Pt(10)
            usage_content = doc.add_paragraph('（在此添加使用说明）')
            for run in usage_content.runs:
                run.font.size = Pt(8)
            doc.add_paragraph()

            # 分隔线（8pt）
            sep_para4 = doc.add_paragraph('─' * 50)
            sep_para4.runs[0].font.size = Pt(8)
            doc.add_paragraph()

            # 注意事项区（10pt 标题，8pt 内容）
            notes_heading = doc.add_heading('注意事项', level=2)
            notes_heading.runs[0].font.size = Pt(10)
            notes_content = doc.add_paragraph('（在此添加注意事项）')
            for run in notes_content.runs:
                run.font.size = Pt(8)

            # 保存文档
            doc.save(str(doc_path))
            logger.info(f"已创建 Word 文档: {doc_path}")

            # 打开文档
            import sys
            if sys.platform == "win32":
                import os
                os.startfile(str(doc_path))
            elif sys.platform == "darwin":
                subprocess.Popen(['open', str(doc_path)])
            else:
                subprocess.Popen(['xdg-open', str(doc_path)])

        except Exception as e:
            logger.warning(f"创建 Word 文档失败: {e}", exc_info=True)

    def _log_performance(self, asset_name: str, package_type, size: int,
                         perf_steps: dict, total_time: float, asset_path: Path = None) -> None:
        """记录性能数据到日志文件

        Args:
            asset_name: 资产名称
            package_type: 资产类型
            size: 资产大小（字节）
            perf_steps: 各步骤耗时字典
            total_time: 总耗时（秒）
            asset_path: 资产路径（用于生成文件结构）
        """
        try:
            from datetime import datetime

            # 性能日志文件路径
            perf_log_path = Path("project-logs/PERFORMANCE.md")

            # 格式化时间戳
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            # 格式化资产大小
            size_str = self._file_ops.format_size(size)

            # 构建性能记录
            perf_record = f"\n### {asset_name}\n\n"
            perf_record += f"- **时间戳**: {timestamp}\n"
            perf_record += f"- **资产类型**: {package_type.display_name}\n"
            perf_record += f"- **资产大小**: {size_str} ({size:,} 字节)\n"
            perf_record += f"- **总耗时**: {total_time:.2f} 秒\n"
            perf_record += f"\n**各步骤耗时**:\n\n"

            for step_name, step_time in perf_steps.items():
                percentage = (step_time / total_time * 100) if total_time > 0 else 0
                perf_record += f"- {step_name}: {step_time:.2f} 秒 ({percentage:.1f}%)\n"

            # 添加文件结构
            if asset_path and asset_path.exists():
                perf_record += f"\n**文件结构**:\n\n"
                perf_record += "```\n"
                perf_record += self._generate_tree_structure(asset_path, max_depth=3)
                perf_record += "```\n"

            perf_record += f"\n---\n"

            # 追加到日志文件
            with open(perf_log_path, 'a', encoding='utf-8') as f:
                f.write(perf_record)

            logger.info(f"性能数据已记录: {asset_name}, 总耗时 {total_time:.2f} 秒")

        except Exception as e:
            logger.warning(f"记录性能数据失败: {e}", exc_info=True)

    def _generate_tree_structure(self, path: Path, prefix: str = "", max_depth: int = 3, 
                                  current_depth: int = 0, max_files: int = 50) -> str:
        """生成目录树结构（限制深度和文件数量）

        Args:
            path: 目录路径
            prefix: 前缀字符串（用于缩进）
            max_depth: 最大深度
            current_depth: 当前深度
            max_files: 最大文件数量

        Returns:
            目录树字符串
        """
        if current_depth >= max_depth:
            return ""

        tree = ""
        file_count = 0

        try:
            items = sorted(path.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))
            
            for i, item in enumerate(items):
                if file_count >= max_files:
                    tree += f"{prefix}└── ... (省略剩余文件)\n"
                    break

                is_last = (i == len(items) - 1)
                connector = "└── " if is_last else "├── "
                
                if item.is_dir():
                    tree += f"{prefix}{connector}{item.name}/\n"
                    
                    # 递归处理子目录
                    if current_depth + 1 < max_depth:
                        extension = "    " if is_last else "│   "
                        tree += self._generate_tree_structure(
                            item, 
                            prefix + extension, 
                            max_depth, 
                            current_depth + 1,
                            max_files - file_count
                        )
                else:
                    # 显示文件大小
                    try:
                        file_size = item.stat().st_size
                        size_str = self._file_ops.format_size(file_size)
                        tree += f"{prefix}{connector}{item.name} ({size_str})\n"
                    except:
                        tree += f"{prefix}{connector}{item.name}\n"
                    
                    file_count += 1

        except PermissionError:
            tree += f"{prefix}└── [权限不足]\n"
        except Exception as e:
            tree += f"{prefix}└── [错误: {e}]\n"

        return tree

    def _format_size(self, size: int) -> str:
        """格式化文件大小 - 委托给 FileOperations"""
        return self._file_ops.format_size(size)
    
    def _set_plugin_default_thumbnail(self, asset: Asset) -> None:
        """为插件类型资产设置默认缩略图
        
        Args:
            asset: 资产对象
        """
        try:
            # 获取默认插件图标路径
            if getattr(sys, 'frozen', False):
                # 打包后的路径
                base_path = Path(sys._MEIPASS)
            else:
                # 开发环境路径
                base_path = Path(__file__).parent.parent.parent.parent
            
            default_icon_path = base_path / "resources" / "icons" / "plugin_default.png"
            
            if not default_icon_path.exists():
                logger.warning(f"默认插件图标不存在: {default_icon_path}")
                return
            
            # 确保缩略图目录存在
            thumbnails_dir = self.thumbnails_dir
            if not thumbnails_dir.exists():
                thumbnails_dir.mkdir(parents=True, exist_ok=True)
            
            # 生成缩略图文件名
            thumbnail_filename = f"{asset.id}.png"
            thumbnail_path = thumbnails_dir / thumbnail_filename
            
            # 复制默认图标作为缩略图
            import shutil
            shutil.copy2(default_icon_path, thumbnail_path)
            
            # 更新资产的缩略图路径
            asset.thumbnail_path = thumbnail_path
            
            logger.info(f"已为插件资产设置默认缩略图: {asset.name}")
            
        except Exception as e:
            logger.warning(f"设置插件默认缩略图失败: {e}", exc_info=True)
    
    def _set_others_default_thumbnail(self, asset: Asset, wrapper_path: Path) -> None:
        """为 OTHERS 类型资产设置默认缩略图
        
        根据资产内容类型（图片/视频/音频）设置不同的默认图标：
        - 图片：使用第一张图片作为缩略图
        - 视频：使用视频默认图标
        - 音频：使用音频默认图标
        - 其他：不设置（保持空白）
        
        Args:
            asset: 资产对象
            wrapper_path: 资产包装文件夹路径
        """
        try:
            others_folder = wrapper_path / "Others"
            if not others_folder.exists():
                return
            
            # 定义文件扩展名
            image_extensions = {'.png', '.jpg', '.jpeg', '.bmp', '.tga', '.tiff', '.tif'}
            video_extensions = {'.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm'}
            audio_extensions = {'.wav', '.mp3', '.ogg', '.flac', '.aac', '.m4a'}
            
            # 扫描 Others 文件夹，检测资产类型
            has_images = False
            has_videos = False
            has_audios = False
            first_image = None
            
            for file in others_folder.rglob('*'):
                if file.is_file():
                    ext = file.suffix.lower()
                    if ext in image_extensions:
                        has_images = True
                        if first_image is None:
                            first_image = file
                    elif ext in video_extensions:
                        has_videos = True
                    elif ext in audio_extensions:
                        has_audios = True
            
            # 确保缩略图目录存在
            thumbnails_dir = self.thumbnails_dir
            if not thumbnails_dir.exists():
                thumbnails_dir.mkdir(parents=True, exist_ok=True)
            
            thumbnail_filename = f"{asset.id}.png"
            thumbnail_path = thumbnails_dir / thumbnail_filename
            
            # 根据资产类型设置缩略图
            if has_images and first_image:
                # 图片资源：使用第一张图片作为缩略图
                try:
                    from PIL import Image
                    img = Image.open(first_image)
                    # 缩放到合适的尺寸（172x115）
                    img.thumbnail((172, 115), Image.Resampling.LANCZOS)
                    # 创建白色背景
                    background = Image.new('RGB', (172, 115), (255, 255, 255))
                    # 居中粘贴
                    offset = ((172 - img.width) // 2, (115 - img.height) // 2)
                    if img.mode == 'RGBA':
                        background.paste(img, offset, img)
                    else:
                        background.paste(img, offset)
                    background.save(thumbnail_path, 'PNG')
                    asset.thumbnail_path = thumbnail_path
                    logger.info(f"已为图片资产设置缩略图: {asset.name}")
                    return
                except Exception as e:
                    logger.warning(f"生成图片缩略图失败: {e}")
            
            # 视频或音频资源：使用默认图标
            if has_videos or has_audios:
                # 获取默认图标路径
                if getattr(sys, 'frozen', False):
                    base_path = Path(sys._MEIPASS)
                else:
                    base_path = Path(__file__).parent.parent.parent.parent
                
                if has_videos:
                    default_icon_path = base_path / "resources" / "icons" / "video_default.png"
                    icon_type = "视频"
                else:
                    default_icon_path = base_path / "resources" / "icons" / "audio_default.png"
                    icon_type = "音频"
                
                if not default_icon_path.exists():
                    logger.warning(f"默认{icon_type}图标不存在: {default_icon_path}")
                    return
                
                # 复制默认图标作为缩略图
                import shutil
                shutil.copy2(default_icon_path, thumbnail_path)
                asset.thumbnail_path = thumbnail_path
                logger.info(f"已为{icon_type}资产设置默认缩略图: {asset.name}")
                return
            
        except Exception as e:
            logger.warning(f"设置 OTHERS 默认缩略图失败: {e}", exc_info=True)
    def _process_others_asset(self, asset_path: Path, inner_folder: Path,
                              original_filename: str, progress_callback: Optional[Callable],
                              perf_steps: dict) -> bool:
        """处理 OTHERS 类型资产：智能提取并分类整理

        Args:
            asset_path: 资产源路径
            inner_folder: 目标文件夹（Others/）
            original_filename: 原始文件名
            progress_callback: 进度回调
            perf_steps: 性能监控字典

        Returns:
            bool: 是否成功
        """
        import zipfile
        import tempfile
        import time
        import re
        from ..utils.archive_extractor import is_ad_file

        # 子步骤 1: 文件名规范化
        sub_step_start = time.time()

        def _normalize_others_folder_name(raw_name: str) -> str:
            name = (raw_name or "").strip()
            if not name:
                return "NewAsset"

            # 去掉文件扩展名
            if '.' in name:
                name = name.rsplit('.', 1)[0]

            # 处理非法字符
            name = re.sub(r'[<>:"/\\|?*]', ' ', name)

            # 递归清理后缀
            while True:
                old_name = name
                name = re.sub(r'\s*[\(（]\s*\d+\s*[\)）]\s*$', '', name, flags=re.IGNORECASE)
                name = re.sub(r'\s*[-_ ]\s*(copy|副本|最终版|final|备份)\s*$', '', name, flags=re.IGNORECASE)
                name = re.sub(r'[-_ ]\d{4}[-_ ]?\d{2}[-_ ]?\d{2}(?:[-_ ]?\d{2,6})?\s*$', '', name)
                name = re.sub(r'\s*[-_ ]\s*v?\d+(?:\.\d+){0,3}\s*$', '', name, flags=re.IGNORECASE)

                if name == old_name:
                    break

            # 清理首尾空格和特殊连字符
            name = re.sub(r'\s+', ' ', name).strip(' ._-')
            return name or "NewAsset"

        raw_subfolder_name = original_filename if original_filename else asset_path.name
        subfolder_name = _normalize_others_folder_name(raw_subfolder_name)
        asset_subfolder = inner_folder / subfolder_name
        asset_subfolder.mkdir(parents=True, exist_ok=True)
        logger.info(f"OTHERS 类型子文件夹: {subfolder_name} (原始: {raw_subfolder_name})")
        perf_steps['OTHERS-01-文件名规范化'] = time.time() - sub_step_start

        # 子步骤 2: 文件扫描和解压
        sub_step_start = time.time()
        model_extensions = {'.fbx', '.obj', '.gltf', '.glb', '.abc', '.usd', '.usda', '.usdc'}
        texture_extensions = {'.png', '.jpg', '.jpeg', '.tga', '.bmp', '.exr', '.hdr', '.tif', '.tiff'}

        models = {}
        textures = {}
        temp_dirs = []
        nested_zip_count = 0
        scanned_files = 0
        skipped_ad_files = 0

        try:
            def scan_and_extract(directory):
                """递归扫描目录，自动解压嵌套 zip"""
                nonlocal nested_zip_count, scanned_files, skipped_ad_files
                for item in directory.iterdir():
                    scanned_files += 1
                    if item.is_file():
                        if is_ad_file(item):
                            skipped_ad_files += 1
                            logger.debug(f"🚫 跳过广告文件: {item.name}")
                            continue

                        ext = item.suffix.lower()
                        if ext == '.zip':
                            nested_zip_count += 1
                            zip_extract_start = time.time()
                            try:
                                nested_temp = Path(tempfile.mkdtemp(prefix='ue_toolkit_nested_'))
                                temp_dirs.append(nested_temp)
                                logger.info(f"📦 解压嵌套压缩包 #{nested_zip_count}: {item.name}")

                                with zipfile.ZipFile(item, 'r') as zip_ref:
                                    file_count = 0
                                    for z_info in zip_ref.infolist():
                                        if z_info.is_dir():
                                            continue

                                        try:
                                            original_name = z_info.filename.encode('cp437').decode('gbk')
                                        except:
                                            original_name = z_info.filename

                                        # 清理文件名中的尾随空格和点
                                        parts = original_name.split('/')
                                        cleaned_parts = []
                                        for part in parts:
                                            cleaned = part.rstrip(' .')
                                            if not cleaned and part:
                                                cleaned = '_'
                                            cleaned_parts.append(cleaned)
                                        cleaned_name = '/'.join(cleaned_parts)

                                        target_path = nested_temp / cleaned_name
                                        target_path.parent.mkdir(parents=True, exist_ok=True)

                                        with zip_ref.open(z_info) as source:
                                            with open(target_path, 'wb') as target:
                                                target.write(source.read())
                                        file_count += 1

                                zip_extract_time = time.time() - zip_extract_start
                                logger.info(f"  ✓ 解压完成: {file_count} 个文件, 耗时 {zip_extract_time:.2f}秒")
                                perf_steps[f'OTHERS-02-解压ZIP#{nested_zip_count}'] = zip_extract_time

                                scan_and_extract(nested_temp)
                            except Exception as e:
                                logger.warning(f"解压嵌套 zip 失败: {item}, {e}")
                        elif ext in model_extensions:
                            if item.name not in models:
                                models[item.name] = (item, item.stat().st_size)
                        elif ext in texture_extensions:
                            current_size = item.stat().st_size
                            if item.name not in textures or current_size > textures[item.name][1]:
                                textures[item.name] = (item, current_size)
                    elif item.is_dir():
                        if item.name.lower() in ['__macosx', '.git']:
                            continue
                        scan_and_extract(item)

            scan_and_extract(asset_path)
            scan_time = time.time() - sub_step_start
            perf_steps['OTHERS-02-文件扫描'] = scan_time

            logger.info(f"📊 扫描完成: 扫描 {scanned_files} 个项目, 跳过广告 {skipped_ad_files} 个")
            logger.info(f"📊 发现: 嵌套ZIP {nested_zip_count} 个, 模型 {len(models)} 个, 贴图 {len(textures)} 个")
            logger.info(f"📊 扫描耗时: {scan_time:.2f}秒")

            # 子步骤 3: 文件复制
            sub_step_start = time.time()
            total_files = len(models) + len(textures)
            total_bytes = sum(size for _, size in models.values()) + \
                          sum(size for _, size in textures.values())

            if total_files == 0:
                logger.error("❌ 未找到任何有用的文件")
                return False

            logger.info(f"📊 准备复制: {total_files} 个文件, 总大小 {self._file_ops.format_size(total_bytes)}")

            copied_bytes = 0
            copied_files = 0

            def copy_file_with_progress(source_file, target_file, file_size, file_type):
                nonlocal copied_bytes, copied_files

                file_copy_start = time.time()
                logger.debug(f"📄 复制 {file_type}: {source_file.name} ({self._file_ops.format_size(file_size)})")

                if file_size < 1024 * 1024:
                    import shutil
                    shutil.copy2(source_file, target_file)
                    copied_bytes += file_size

                    if progress_callback:
                        progress = 10 + int(copied_bytes / total_bytes * 70)
                        progress_callback(progress, 100, f"复制{file_type}: {source_file.name}")
                else:
                    chunk_size = 1024 * 1024

                    with open(source_file, 'rb') as fsrc:
                        with open(target_file, 'wb') as fdst:
                            while True:
                                chunk = fsrc.read(chunk_size)
                                if not chunk:
                                    break
                                fdst.write(chunk)
                                copied_bytes += len(chunk)

                                if progress_callback:
                                    progress = 10 + int(copied_bytes / total_bytes * 70)
                                    progress_callback(progress, 100, f"复制{file_type}: {source_file.name}")

                    import shutil
                    shutil.copystat(str(source_file), str(target_file))

                copied_files += 1
                file_copy_time = time.time() - file_copy_start
                logger.debug(f"  ✓ 复制完成: {file_copy_time:.2f}秒")

            # 复制模型文件
            if models:
                mesh_folder = asset_subfolder / "Mesh"
                mesh_folder.mkdir(parents=True, exist_ok=True)
                logger.info(f"📁 复制模型文件到: {mesh_folder}")
                for filename, (source_file, file_size) in models.items():
                    target_file = mesh_folder / filename
                    try:
                        copy_file_with_progress(source_file, target_file, file_size, "模型")
                    except Exception as e:
                        logger.error(f"❌ 复制模型失败: {source_file} -> {target_file}, {e}")
                        return False

            # 复制贴图文件
            if textures:
                textures_folder = asset_subfolder / "Textures"
                textures_folder.mkdir(parents=True, exist_ok=True)
                logger.info(f"📁 复制贴图文件到: {textures_folder}")
                for filename, (source_file, file_size) in textures.items():
                    target_file = textures_folder / filename
                    try:
                        copy_file_with_progress(source_file, target_file, file_size, "贴图")
                    except Exception as e:
                        logger.error(f"❌ 复制贴图失败: {source_file} -> {target_file}, {e}")
                        return False

            copy_time = time.time() - sub_step_start
            perf_steps['OTHERS-03-文件复制'] = copy_time
            logger.info(f"✅ 复制完成: {copied_files}/{total_files} 个文件, 耗时 {copy_time:.2f}秒")
            logger.info(f"📊 平均速度: {self._file_ops.format_size(total_bytes / copy_time)}/s")

        finally:
            # 子步骤 4: 临时目录清理
            sub_step_start = time.time()
            import time as time_module
            import gc

            gc.collect()
            time_module.sleep(0.5)

            cleaned_count = 0
            failed_count = 0
            for temp_dir in temp_dirs:
                try:
                    import shutil
                    for attempt in range(3):
                        try:
                            shutil.rmtree(temp_dir)
                            cleaned_count += 1
                            break
                        except PermissionError:
                            if attempt < 2:
                                logger.debug(f"临时目录被占用，重试 ({attempt + 1}/3): {temp_dir}")
                                time_module.sleep(1)
                            else:
                                failed_count += 1
                                logger.warning(f"临时目录清理失败（文件被占用）: {temp_dir}")
                        except Exception as e:
                            failed_count += 1
                            logger.warning(f"清理临时目录失败: {temp_dir}, {e}")
                            break
                except Exception as e:
                    failed_count += 1
                    logger.warning(f"清理临时目录时出错: {temp_dir}, {e}")

            cleanup_time = time.time() - sub_step_start
            perf_steps['OTHERS-04-临时清理'] = cleanup_time
            logger.info(f"🧹 清理完成: {cleaned_count}/{len(temp_dirs)} 个临时目录, 失败 {failed_count} 个, 耗时 {cleanup_time:.2f}秒")

        return True



    def remove_asset(self, asset_id: str, delete_physical: bool = False) -> bool:
        """删除资产"""
        try:
            asset = self.get_asset(asset_id)
            if not asset:
                error_msg = f"未找到资产: {asset_id}"
                logger.error(error_msg)
                self.error_occurred.emit(error_msg)
                return False

            # 删除物理文件
            if delete_physical and asset.path and asset.path.exists():
                try:
                    if asset.path.is_dir():
                        success = self._file_ops.safe_remove_tree(asset.path)
                    else:
                        success = self._file_ops.safe_remove_file(asset.path)
                    if not success:
                        raise Exception("文件删除操作返回 False")
                    logger.info(f"已删除资产文件: {asset.path}")
                except Exception as e:
                    error_msg = f"删除物理文件失败: {e}"
                    logger.error(error_msg, exc_info=True)
                    self.error_occurred.emit(error_msg)
                    return False

            # 从列表中删除（就地修改，保持与 _asset_core.assets 的引用一致）
            self.assets[:] = [a for a in self.assets if a.id != asset_id]

            # 删除缩略图
            if asset.thumbnail_path and asset.thumbnail_path.exists():
                try:
                    asset.thumbnail_path.unlink()
                except Exception as e:
                    logger.warning(f"删除缩略图失败: {e}")

            # 删除关联文档（支持 .docx 和旧的 .txt/.md 格式）
            if self.documents_dir:
                # 优先删除 .docx 文件
                doc_path = self.documents_dir / f"{asset_id}.docx"
                if doc_path.exists():
                    try:
                        doc_path.unlink()
                        logger.info(f"已删除资产文档: {doc_path}")
                    except Exception as e:
                        logger.warning(f"删除关联文档失败: {e}")
                
                # 兼容旧格式：删除 .txt 文件
                txt_path = self.documents_dir / f"{asset_id}.txt"
                if txt_path.exists():
                    try:
                        txt_path.unlink()
                        logger.info(f"已删除旧格式文档: {txt_path}")
                    except Exception as e:
                        logger.warning(f"删除旧格式文档失败: {e}")
                
                # 兼容旧格式：删除 .md 文件
                md_path = self.documents_dir / f"{asset_id}.md"
                if md_path.exists():
                    try:
                        md_path.unlink()
                        logger.info(f"已删除旧格式文档: {md_path}")
                    except Exception as e:
                        logger.warning(f"删除旧格式文档失败: {e}")

            self._save_config()

            logger.info(f"删除资产成功: {asset.name} (物理删除: {delete_physical})")
            self.asset_removed.emit(asset_id)

            return True

        except Exception as e:
            error_msg = f"删除资产失败: {e}"
            logger.error(error_msg, exc_info=True)
            self.error_occurred.emit(error_msg)
            return False

    def get_asset(self, asset_id: str) -> Optional[Asset]:
        """获取指定资产 - 委托给 AssetCore"""
        return self._asset_core.get_asset(asset_id)

    def get_all_assets(self, category: Optional[str] = None) -> List[Asset]:
        """获取所有资产 - 委托给 AssetCore"""
        return self._asset_core.get_all_assets(category)

    def get_all_categories(self) -> List[str]:
        """获取所有分类 - 委托给 AssetCore"""
        return self._asset_core.get_all_categories()

    def get_all_asset_names(self) -> List[str]:
        """获取所有资产名称"""
        return [asset.name for asset in self.assets]


    def add_category(self, category_name: str) -> bool:
        """添加新分类 - 委托给 AssetCore + 文件系统操作"""
        if not self._asset_core.add_category(category_name):
            return False

        library_path = self.get_asset_library_path()
        if library_path and library_path.exists():
            category_folder = library_path / category_name.strip()
            if not category_folder.exists():
                try:
                    category_folder.mkdir(parents=True, exist_ok=True)
                except Exception as e:
                    logger.error(f"创建分类文件夹失败: {e}", exc_info=True)

        self._save_config()
        logger.info(f"已添加新分类: {category_name}")
        return True

    def remove_category(self, category_name: str) -> bool:
        """删除分类"""
        if category_name == "默认分类":
            logger.warning("不能删除默认分类")
            return False

        if category_name not in self.categories:
            return False

        assets_in_category = [a for a in self.assets if a.category == category_name]
        if assets_in_category:
            library_path = self.get_asset_library_path()
            if library_path and library_path.exists():
                for asset in assets_in_category:
                    self._move_asset_to_category(asset, "默认分类")
            else:
                for asset in assets_in_category:
                    asset.category = "默认分类"
            logger.info(f"已将 {len(assets_in_category)} 个资产从 {category_name} 移至默认分类")

        self.categories.remove(category_name)

        library_path = self.get_asset_library_path()
        if library_path and library_path.exists():
            category_folder = library_path / category_name
            if category_folder.exists():
                try:
                    if not any(category_folder.iterdir()):
                        category_folder.rmdir()
                except Exception as e:
                    logger.error(f"删除分类文件夹失败: {e}", exc_info=True)

        self._save_config()
        logger.info(f"已删除分类: {category_name}")
        return True

    # ─── 搜索和排序（委托给 SearchEngine）──────────────────

    def search_assets(self, search_text: str,
                      category: Optional[str] = None) -> List[Asset]:
        """搜索资产 - 委托给 SearchEngine"""
        if not search_text or not search_text.strip():
            return self.get_all_assets(category)

        candidates = self.get_all_assets(category)
        return self._search_engine.search(candidates, search_text, category=None)

    def sort_assets(self, assets: List[Asset], sort_method: str) -> List[Asset]:
        """排序资产 - 委托给 SearchEngine"""
        if not assets:
            return []

        try:
            # 保持原有排序逻辑以确保向后兼容
            if sort_method in ["添加时间（最新）", "🕐 最新添加"]:
                return sorted(assets,
                              key=lambda x: x.created_time if x.created_time else datetime.min,
                              reverse=True)
            elif sort_method in ["添加时间（最早）", "🕐 最早添加"]:
                return sorted(assets,
                              key=lambda x: x.created_time if x.created_time else datetime.min)
            elif sort_method in ["名称（A-Z）", "🔤 名称 A-Z"]:
                return sorted(assets, key=lambda x: x.name.lower())
            elif sort_method in ["名称（Z-A）", "🔤 名称 Z-A"]:
                return sorted(assets, key=lambda x: x.name.lower(), reverse=True)
            elif sort_method in ["分类（A-Z）", "📁 分类 A-Z"]:
                return sorted(assets, key=lambda x: (x.category.lower(), x.name.lower()))
            elif sort_method in ["分类（Z-A）", "📁 分类 Z-A"]:
                return sorted(assets, key=lambda x: (x.category.lower(), x.name.lower()),
                              reverse=True)
            else:
                return sorted(assets, key=lambda x: x.created_time, reverse=True)
        except Exception as e:
            logger.error(f"排序资产时出错: {e}", exc_info=True)
            return assets

    # ─── 资产信息更新 ──────────────────────────────────────

    def update_asset_description(self, asset_id: str, description: str) -> bool:
        """更新资产描述 - 委托给 AssetCore"""
        if not self._asset_core.update_asset(asset_id, description=description):
            return False

        self._save_config()
        asset = self._asset_core.get_asset(asset_id)
        logger.info(f"资产描述已更新并保存: {asset.name if asset else asset_id}")
        return True

    def update_asset_info(self, asset_id: str, new_name: Optional[str] = None,
                          new_category: Optional[str] = None) -> bool:
        """更新资产信息（名称和/或分类）"""
        asset = self._asset_core.get_asset(asset_id)
        if not asset:
            logger.warning(f"资产不存在，无法更新信息: {asset_id}")
            return False

        old_name = asset.name
        old_category = asset.category
        old_path = asset.path

        # 更新名称（同时重命名文件夹）
        if new_name is not None and new_name.strip() and new_name != old_name:
            new_name = new_name.strip()
            if not self._rename_asset_folder(asset, new_name):
                logger.warning(f"资产文件夹重命名失败，仅更新配置: {old_name} -> {new_name}")
            # 更新配置中的名称和路径
            self._asset_core.update_asset(asset_id, name=new_name, path=str(asset.path))

        # 更新分类（移动到新分类文件夹）
        if new_category is not None and new_category.strip() and new_category != old_category:
            new_category = new_category.strip()
            if new_category not in self.categories:
                self.add_category(new_category)

            if not self._move_asset_to_category(asset, new_category):
                # 移动失败，仅更新配置中的分类
                self._asset_core.update_asset(asset_id, category=new_category)
                logger.warning(f"资产物理移动失败，仅更新配置: {old_category} -> {new_category}")
            else:
                # 移动成功，更新配置中的分类和路径
                self._asset_core.update_asset(asset_id, category=new_category, path=str(asset.path))

        # 重建该资产的拼音缓存
        self._search_engine.build_pinyin_cache([asset])

        self._save_config()

        changes = []
        if new_name and new_name != old_name:
            changes.append(f"名称: {old_name} -> {new_name}")
        if new_category and new_category != old_category:
            changes.append(f"分类: {old_category} -> {new_category}")
        if changes:
            logger.info(f"资产信息已更新: {', '.join(changes)}")

        return True

    def _calculate_size(self, path: Path) -> int:
        """计算文件或文件夹大小 - 委托给 FileOperations"""
        return self._file_ops.calculate_size(path)


    # ─── 资产库路径管理 ────────────────────────────────────

    def get_asset_library_path(self) -> Optional[Path]:
        """获取资产库路径"""
        config = self.config_manager.load_user_config()
        asset_library_path = (config.get("current_asset_library", "")
                               or config.get("asset_library_path", ""))
        if not asset_library_path:
            for lib in config.get("asset_libraries", []):
                p = lib.get("path", "")
                if p and Path(p).exists():
                    asset_library_path = p
                    break
        if not asset_library_path:
            for key in config.get("asset_library_configs", {}).keys():
                if Path(key).exists():
                    asset_library_path = key
                    break

        if asset_library_path:
            return Path(asset_library_path)
        return None

    def set_asset_library_path(self, library_path: Path) -> bool:
        """设置资产库路径"""
        try:
            if not library_path.exists():
                logger.info(f"资产库路径不存在，正在创建: {library_path}")
                library_path.mkdir(parents=True, exist_ok=True)

            current_lib_path = self.get_asset_library_path()
            if current_lib_path:
                logger.info(f"保存当前资产库配置: {current_lib_path}")
                self._save_config()

            logger.info(f"正在加载配置文件...")
            config = self.config_manager.load_user_config() or {}
            if "_version" not in config:
                config["_version"] = "2.0.0"

            config["current_asset_library"] = str(library_path)
            libs = config.setdefault("asset_libraries", [])
            if not any(l.get("path") == str(library_path) for l in libs):
                libs.append({"path": str(library_path), "name": "主资产库", "last_opened": ""})
            logger.info(f"正在保存资产库路径到配置: {library_path}")
            
            save_result = self.config_manager.save_user_config(config, backup_reason="set_library_path")
            
            if not save_result:
                logger.error(f"❌ 保存资产库路径失败！配置未写入文件")
                return False

            logger.info(f"✅ 资产库路径已成功保存到配置文件")
            
            logger.info(f"正在重新加载配置...")
            self._load_config()

            logger.info(f"✅ 资产库路径已切换至: {library_path}")
            return True

        except Exception as e:
            logger.error(f"❌ 设置资产库路径失败: {e}", exc_info=True)
            return False

    def _rename_asset_folder(self, asset: Asset, new_name: str) -> bool:
        """重命名资产文件夹
        
        Args:
            asset: 资产对象
            new_name: 新的资产名称
            
        Returns:
            bool: 是否成功重命名
        """
        try:
            old_path = asset.path
            if not old_path.exists():
                logger.warning(f"资产文件夹不存在，无法重命名: {old_path}")
                return False
            
            # 构建新路径（保持在同一分类文件夹下）
            new_path = old_path.parent / new_name
            
            # 检查新路径是否已存在
            if new_path.exists():
                logger.warning(f"目标文件夹已存在，无法重命名: {new_path}")
                return False
            
            # 重命名文件夹
            old_path.rename(new_path)
            logger.info(f"重命名资产文件夹: {old_path.name} -> {new_name}")
            
            # 更新资产路径
            asset.path = new_path
            return True
            
        except Exception as e:
            logger.error(f"重命名资产文件夹失败: {e}", exc_info=True)
            return False
    
    def _move_asset_to_category(self, asset: Asset, new_category: str) -> bool:
        """将资产物理移动到新分类文件夹"""
        library_path = self.get_asset_library_path()
        if not library_path or not library_path.exists():
            logger.warning("资产库路径未设置，无法移动资产")
            return False

        try:
            new_category_folder = library_path / new_category
            if not new_category_folder.exists():
                new_category_folder.mkdir(parents=True, exist_ok=True)

            old_path = asset.path
            new_path = new_category_folder / old_path.name

            if old_path.exists():
                shutil.move(str(old_path), str(new_path))
                logger.info(f"移动资产: {old_path} -> {new_path}")
                asset.path = new_path

            asset.category = new_category
            return True

        except Exception as e:
            logger.error(f"移动资产失败: {e}", exc_info=True)
            return False

    def _sync_category_folders(self):
        """同步分类文件夹到资产库"""
        library_path = self.get_asset_library_path()
        if not library_path:
            return

        try:
            for category in self.categories:
                category_folder = library_path / category
                if not category_folder.exists():
                    category_folder.mkdir(parents=True, exist_ok=True)
            logger.info("分类文件夹同步完成")
        except Exception as e:
            logger.error(f"同步分类文件夹失败: {e}", exc_info=True)

    # ─── 预览相关（委托给 AssetPreviewCoordinator）─────────

    def set_preview_project(self, project_path: Path) -> bool:
        """设置预览工程路径"""
        return self._preview_coordinator.set_preview_project(project_path)

    def get_preview_project(self) -> Optional[Path]:
        """获取预览工程路径"""
        return self._preview_coordinator.get_preview_project()

    def get_additional_preview_projects(self) -> List[str]:
        """获取额外的预览工程路径列表"""
        return self._preview_coordinator.get_additional_preview_projects()

    def get_additional_preview_projects_with_names(self) -> List[Dict[str, str]]:
        """获取额外的预览工程路径和名称列表"""
        return self._preview_coordinator.get_additional_preview_projects_with_names()

    def set_additional_preview_projects(self, project_paths: List[str]) -> bool:
        """设置额外的预览工程路径列表"""
        return self._preview_coordinator.set_additional_preview_projects(project_paths)

    def set_additional_preview_projects_with_names(self,
                                                    projects: List[Dict[str, str]]) -> bool:
        """设置额外的预览工程路径和名称列表"""
        return self._preview_coordinator.set_additional_preview_projects_with_names(projects)

    def get_use_symlink_preview(self) -> bool:
        """获取是否使用符号链接进行预览（已废除，始终返回 False）。"""
        return False

    def set_use_symlink_preview(self, enabled: bool) -> bool:
        """设置是否使用符号链接进行预览（已废除，不执行任何操作）。"""
        return True

    def clean_preview_project(self) -> bool:
        """清理预览工程的Content文件夹"""
        return self._preview_coordinator.clean_preview_project(
            error_callback=lambda msg: self.error_occurred.emit(msg)
        )

    def preview_asset(self, asset_id: str, progress_callback=None,
                      preview_project_path: Optional[Path] = None) -> bool:
        """预览资产"""
        try:
            asset = self.get_asset(asset_id)
            if not asset:
                error_msg = f"未找到资产: {asset_id}"
                logger.error(error_msg)
                self.error_occurred.emit(error_msg)
                return False

            if preview_project_path is None:
                preview_project = self._preview_coordinator.get_preview_project()
            else:
                preview_project = preview_project_path

            if not preview_project or not preview_project.exists():
                error_msg = "预览工程未设置或不存在，请先设置预览工程"
                logger.error(error_msg)
                self.error_occurred.emit(error_msg)
                return False

            # 关闭当前正在运行的预览工程
            running_process = self._preview_coordinator.check_preview_project_running(
                preview_project
            )
            if running_process or self.current_preview_process or self.current_preview_project_path:
                logger.info("检测到正在运行的预览工程，准备关闭...")
                if progress_callback:
                    progress_callback(0, 1, "正在关闭之前的预览工程...")

                # 确保 coordinator 有正确的进程引用
                if running_process:
                    self._preview_coordinator.current_preview_process = running_process
                    self._preview_coordinator.current_preview_project_path = preview_project

                if not self._preview_coordinator.close_current_preview_if_running():
                    # 关闭失败时不直接报错，尝试继续（UE 可能已经在关闭中）
                    logger.warning("自动关闭预览工程失败，尝试继续预览...")
                    import time
                    time.sleep(2)
                    # 再检查一次
                    still_running = self._preview_coordinator.check_preview_project_running(preview_project)
                    if still_running:
                        error_msg = "无法关闭当前正在运行的预览工程，请手动关闭后重试"
                        logger.error(error_msg)
                        self.error_occurred.emit(error_msg)
                        return False
                    else:
                        logger.info("预览工程已自行关闭，继续预览")

            # 在后台线程中执行预览
            thread = threading.Thread(
                target=self._preview_coordinator.do_preview_asset,
                args=(asset, preview_project),
                kwargs={
                    'progress_callback': progress_callback,
                    'on_preview_finished': lambda: self.preview_finished.emit(),
                    'on_error': lambda msg: self.error_occurred.emit(msg),
                    'on_thumbnail_updated': lambda aid, tp: self.thumbnail_updated.emit(aid, tp),
                    'save_config_callback': self._save_config,
                    'thumbnails_dir': self.thumbnails_dir,
                },
                daemon=True
            )
            thread.start()

            self.preview_started.emit(asset_id)
            return True

        except Exception as e:
            error_msg = f"启动预览失败: {e}"
            logger.error(error_msg, exc_info=True)
            self.error_occurred.emit(error_msg)
            return False

    def _do_preview_asset(self, asset: Asset, preview_project: Path,
                           progress_callback=None) -> None:
        """执行资产预览（后台线程）- 委托给 AssetPreviewCoordinator"""
        self._preview_coordinator.do_preview_asset(
            asset=asset,
            preview_project=preview_project,
            progress_callback=progress_callback,
            on_preview_finished=lambda: self.preview_finished.emit(),
            on_error=lambda msg: self.error_occurred.emit(msg),
            on_thumbnail_updated=lambda aid, tp: self.thumbnail_updated.emit(aid, tp),
            save_config_callback=self._save_config,
            thumbnails_dir=self.thumbnails_dir,
        )

    def _launch_unreal_project(self, project_path: Path):
        """委托给 AssetPreviewCoordinator"""
        return self._preview_coordinator._launch_unreal_project(project_path)

    def _check_preview_project_running(self, preview_project: Path):
        """委托给 AssetPreviewCoordinator"""
        return self._preview_coordinator.check_preview_project_running(preview_project)

    def _find_ue_process(self):
        """委托给 AssetPreviewCoordinator"""
        return self._preview_coordinator._find_ue_process()

    def _close_current_preview_if_running(self):
        """委托给 AssetPreviewCoordinator"""
        return self._preview_coordinator.close_current_preview_if_running()

    def migrate_asset(self, asset_id: str, target_project: Path,
                      progress_callback=None) -> bool:
        """将资产迁移到目标工程"""
        try:
            asset = self.get_asset(asset_id)
            if not asset:
                error_msg = f"未找到资产: {asset_id}"
                logger.error(error_msg)
                self.error_occurred.emit(error_msg)
                return False

            return self._preview_coordinator.migrate_asset(
                asset=asset,
                target_project=target_project,
                progress_callback=progress_callback,
                error_callback=lambda msg: self.error_occurred.emit(msg)
            )

        except Exception as e:
            error_msg = f"迁移资产失败: {e}"
            logger.error(error_msg, exc_info=True)
            self.error_occurred.emit(error_msg)
            return False

    def process_screenshot(self, asset_id: str, preview_project: Path) -> None:
        """处理预览截图"""
        try:
            asset = self.get_asset(asset_id)
            if not asset:
                logger.warning(f"未找到资产: {asset_id}")
                return

            self._preview_coordinator.process_screenshot(
                asset=asset,
                preview_project=preview_project,
                thumbnails_dir=self.thumbnails_dir,
                on_thumbnail_updated=lambda aid, tp: self.thumbnail_updated.emit(aid, tp),
                save_config_callback=self._save_config
            )

        except Exception as e:
            logger.error(f"处理截图时出错: {e}", exc_info=True)

    # ─── 向后兼容的拼音方法（委托给 SearchEngine）──────────

    def _get_pinyin(self, text: str) -> str:
        """委托给 SearchEngine"""
        return self._search_engine.get_pinyin(text)

    def _get_pinyin_initials(self, text: str) -> str:
        """委托给 SearchEngine"""
        return self._search_engine.get_pinyin_initials(text)

    def _build_pinyin_cache(self) -> None:
        """委托给 SearchEngine"""
        self._search_engine.build_pinyin_cache(self.assets)

    def _get_asset_pinyin(self, asset_id: str) -> Dict[str, str]:
        """委托给 SearchEngine"""
        return self._search_engine._pinyin_cache.get(asset_id, {
            'name_pinyin': '',
            'name_initials': '',
            'desc_pinyin': '',
            'desc_initials': '',
            'category_pinyin': '',
            'category_initials': ''
        })

    # ─── 向后兼容：_pinyin_cache 属性 ─────────────────────

    @property
    def _pinyin_cache(self) -> Dict[str, Dict[str, str]]:
        """向后兼容：访问 SearchEngine 的拼音缓存"""
        return self._search_engine._pinyin_cache

    @_pinyin_cache.setter
    def _pinyin_cache(self, value: Dict[str, Dict[str, str]]):
        """向后兼容：设置 SearchEngine 的拼音缓存"""
        self._search_engine._pinyin_cache = value

    # ─── 向后兼容：_load_assets_from_config ────────────────

    def _load_assets_from_config(self, assets_data: List[Dict[str, Any]]) -> None:
        """从配置数据加载资产列表"""
        self.assets.clear()

        for asset_data in assets_data:
            try:
                created_time_str = asset_data.get("created_time")
                if created_time_str:
                    try:
                        created_time = datetime.fromisoformat(created_time_str)
                    except (ValueError, TypeError):
                        created_time = datetime.now()
                else:
                    created_time = datetime.now()

                # 解析 package_type（向后兼容：旧数据无此字段默认为 content）
                pkg_type_str = asset_data.get("package_type", "content")
                try:
                    package_type = PackageType(pkg_type_str)
                except (ValueError, KeyError):
                    package_type = PackageType.CONTENT

                asset = Asset(
                    id=asset_data["id"],
                    name=asset_data["name"],
                    asset_type=AssetType(asset_data["asset_type"]),
                    path=Path(asset_data["path"]),
                    category=asset_data.get("category", "默认分类"),
                    package_type=package_type,
                    file_extension=asset_data.get("file_extension", ""),
                    thumbnail_path=Path(asset_data["thumbnail_path"]) if asset_data.get("thumbnail_path") else None,
                    thumbnail_source=asset_data.get("thumbnail_source"),
                    size=asset_data.get("size", 0),
                    created_time=created_time,
                    description=asset_data.get("description", ""),
                    engine_min_version=asset_data.get("engine_min_version", ""),
                    project_file=asset_data.get("project_file", "")
                )

                if asset.path.exists():
                    self.assets.append(asset)
                else:
                    logger.warning(f"资产路径不存在，跳过: {asset.path}")

            except Exception as e:
                logger.error(f"加载资产数据失败: {e}", exc_info=True)

        logger.info(f"已加载 {len(self.assets)} 个资产")
        self.assets_loaded.emit(self.assets)

    # ─── 向后兼容：_migrate_local_config ───────────────────

    def _migrate_local_config(self, config: Dict[str, Any]) -> Dict[str, Any]:
        """委托给 AssetLocalConfigManager"""
        return self._local_config._migrate_local_config(config)

    # ─── 资产自动检测 ──────────────────────────────────────

    def start_auto_detect(self) -> bool:
        """启动资产自动检测（监控资产库 Content 文件夹）"""
        # 自动检测功能已移除
        return False

    def stop_auto_detect(self):
        """停止资产自动检测"""
        # 自动检测功能已移除
        pass

    def refresh_auto_detect(self):
        """刷新 Content 目录快照（资产增删后调用）"""
        # 自动检测功能已移除
        pass
