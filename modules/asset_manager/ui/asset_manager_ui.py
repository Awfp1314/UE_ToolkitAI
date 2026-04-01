# -*- coding: utf-8 -*-

"""
资产管理器 UI
"""

from pathlib import Path
from collections import OrderedDict

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QScrollArea,
    QGridLayout, QLineEdit, QPushButton, QApplication, QDialog, QComboBox,
    QLabel, QFileDialog
)
from PyQt6.QtCore import Qt, QTimer, pyqtSignal, QObject, QUrl, QThread, QPoint, QPoint
from PyQt6.QtGui import QPixmap, QDragEnterEvent, QDropEvent

from core.logger import get_logger
from ui.thumbnail_cache import ThumbnailCache
from modules.base_module_widget import BaseModuleWidget
from modules.asset_manager.logic.asset_controller import AssetController
from .modern_asset_card import ModernAssetCard, CompactAssetCard, format_size, format_time
from .project_selector_dialog import ProjectSelectorDialog
from .add_asset_dialog import AddAssetDialog
from .confirm_dialog import ConfirmDialog
from .asset_detail_dialog import AssetDetailDialog
from ..utils.archive_extractor import ARCHIVE_EXTENSIONS, ArchiveExtractor
from ..utils.asset_structure_analyzer import AssetStructureAnalyzer, StructureType
from ..utils.ue_version_detector import UEVersionDetector
from ..logic.asset_model import AssetType, PackageType

logger = get_logger(__name__)


class _StageSpinnerLabel(QLabel):
    """阶段加载动画"""
    _FRAMES = ["◐", "◓", "◑", "◒"]

    def __init__(self, parent=None):
        super().__init__(parent)
        self._idx = 0
        self.setText(self._FRAMES[0])
        self._timer = QTimer(self)
        self._timer.timeout.connect(self._tick)
        self._timer.start(120)

    def _tick(self):
        self._idx = (self._idx + 1) % len(self._FRAMES)
        self.setText(self._FRAMES[self._idx])

    def stop(self):
        self._timer.stop()


class AddAssetSourceTypeDialog(QDialog):
    """添加资产来源格式选择弹窗"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.selected_source = None  # "archive" | "folder"
        self.drag_position = QPoint()

        self.setModal(True)
        self.setFixedSize(340, 220)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setObjectName("ConfirmDialog")

        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 22, 24, 20)
        layout.setSpacing(10)

        title = QLabel("选择来源格式")
        title.setObjectName("ConfirmDialogMessage")
        layout.addWidget(title)

        hint = QLabel("请选择要添加的资产来源类型")
        hint.setObjectName("ConfirmDialogDetails")
        layout.addWidget(hint)

        layout.addSpacing(6)

        # 两个主按钮并排
        btn_row = QHBoxLayout()
        btn_row.setSpacing(10)

        archive_btn = QPushButton("压缩包")
        archive_btn.setObjectName("ConfirmDialogOkButton")
        archive_btn.setFixedHeight(34)
        archive_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        archive_btn.clicked.connect(lambda: self._select("archive"))
        btn_row.addWidget(archive_btn)

        folder_btn = QPushButton("文件夹")
        folder_btn.setObjectName("ConfirmDialogOkButton")
        folder_btn.setFixedHeight(34)
        folder_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        folder_btn.clicked.connect(lambda: self._select("folder"))
        btn_row.addWidget(folder_btn)

        layout.addLayout(btn_row)

        # 取消按钮单独右对齐
        cancel_row = QHBoxLayout()
        cancel_row.addStretch()
        cancel_btn = QPushButton("取消")
        cancel_btn.setObjectName("ConfirmDialogCancelButton")
        cancel_btn.setFixedSize(72, 28)
        cancel_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        cancel_btn.clicked.connect(self.reject)
        cancel_row.addWidget(cancel_btn)
        layout.addLayout(cancel_row)

    def _select(self, source_type: str):
        self.selected_source = source_type
        self.accept()

    def showEvent(self, event):
        super().showEvent(event)
        self._center_on_parent()

    def _center_on_parent(self):
        """左右居中，上下偏上（1/4 处），相对右侧 RightPanel 区域"""
        self.adjustSize()
        host = self._find_right_panel()
        if host:
            host_global = host.mapToGlobal(host.rect().topLeft())
            x = host_global.x() + (host.width() - self.width()) // 2
            y = host_global.y() + (host.height() - self.height()) // 4
            self.move(x, y)
            return
        screen = QApplication.primaryScreen()
        if screen:
            available = screen.availableGeometry()
            self.move(
                available.x() + (available.width() - self.width()) // 2,
                available.y() + (available.height() - self.height()) // 4,
            )

    def _find_right_panel(self):
        """向上遍历 parent 链，找到 objectName == 'RightPanel' 的组件"""
        w = self.parent()
        while w is not None:
            if hasattr(w, 'objectName') and w.objectName() == "RightPanel":
                return w
            w = w.parent() if hasattr(w, 'parent') and callable(w.parent) else None
        return self.parent()


class AddAssetAnalysisThread(QThread):
    """资产来源识别与分析线程"""

    stage_running = pyqtSignal(int, str)
    stage_done = pyqtSignal(int)
    analysis_done = pyqtSignal(dict)
    analysis_failed = pyqtSignal(str)

    def __init__(self, source_path: Path, source_type: str, password: str = "", parent=None):
        super().__init__(parent)
        self.source_path = source_path
        self.source_type = source_type  # archive | folder
        self.password = password

    def run(self):
        try:
            if self.source_type == "archive":
                self._analyze_archive()
            else:
                self._analyze_folder()
        except Exception as e:
            logger.error(f"资产分析线程异常: {e}", exc_info=True)
            self.analysis_failed.emit(f"分析失败: {e}")

    def _map_package_type(self, structure_type: str) -> PackageType:
        pkg_map = {
            StructureType.CONTENT_PACKAGE.value: PackageType.CONTENT,
            StructureType.UE_PROJECT.value: PackageType.PROJECT,
            StructureType.UE_PLUGIN.value: PackageType.PLUGIN,
            StructureType.LOOSE_ASSETS.value: PackageType.OTHERS,
            StructureType.RAW_3D_FILES.value: PackageType.OTHERS,
            StructureType.MIXED_FILES.value: PackageType.OTHERS,
            StructureType.UNKNOWN.value: PackageType.OTHERS,
        }
        return pkg_map.get(structure_type, PackageType.OTHERS)

    def _analyze_archive(self):
        self.stage_running.emit(0, "正在分析资产版本")

        extractor = ArchiveExtractor()
        try:
            temp_dir = extractor.extract(self.source_path, password=self.password if self.password else None)
        except RuntimeError as e:
            self.analysis_failed.emit(str(e))
            return

        analyzer = AssetStructureAnalyzer()
        analysis = analyzer.analyze(temp_dir)
        self.stage_done.emit(0)

        import time as _time
        _time.sleep(0.4)  # 等待阶段0完成动画播完（0→50%）
        self.stage_running.emit(1, "正在分析资产类型")
        _time.sleep(0.4)  # 等待阶段1开始动画播完（50→75%）
        package_type = self._map_package_type(analysis.structure_type.value)

        engine_version = analysis.engine_version or ""
        if not engine_version:
            detector = UEVersionDetector(logger)
            detect_path = analysis.content_root or analysis.asset_root or temp_dir
            engine_version = detector.detect_asset_min_version(detect_path) or ""

        if package_type == PackageType.PLUGIN and analysis.asset_root:
            resolved_path = analysis.asset_root
            plugin_folder_name = analysis.asset_root.name
        else:
            resolved_path = analysis.asset_root or temp_dir
            plugin_folder_name = ""

        self.stage_done.emit(1)
        self.analysis_done.emit({
            "is_archive": True,
            "package_type": package_type,
            "engine_version": engine_version,
            "suggested_name": analysis.suggested_name,
            "resolved_asset_path": str(resolved_path),
            "archive_content_path": str(resolved_path),
            "archive_temp_dir": str(temp_dir),
            "archive_extractor": extractor,
            "plugin_folder_name": plugin_folder_name,
        })

    def _analyze_folder(self):
        self.stage_running.emit(0, "正在分析资产版本")

        analyzer = AssetStructureAnalyzer()
        analysis = analyzer.analyze(self.source_path)
        package_type = self._map_package_type(analysis.structure_type.value)
        self.stage_done.emit(0)

        import time as _time
        _time.sleep(0.4)  # 等待阶段0完成动画播完（0→50%）
        self.stage_running.emit(1, "正在分析资产类型")
        _time.sleep(0.4)  # 等待阶段1开始动画播完（50→75%）
        engine_version = analysis.engine_version or ""
        if not engine_version:
            detector = UEVersionDetector(logger)
            detect_path = analysis.content_root or analysis.asset_root or self.source_path
            engine_version = detector.detect_asset_min_version(detect_path) or ""

        if package_type in (PackageType.PLUGIN, PackageType.PROJECT) and analysis.asset_root:
            resolved_path = analysis.asset_root
        elif analysis.content_root:
            resolved_path = analysis.content_root
        elif analysis.asset_root:
            resolved_path = analysis.asset_root
        else:
            resolved_path = self.source_path

        plugin_folder_name = resolved_path.name if package_type == PackageType.PLUGIN else ""

        self.stage_done.emit(1)
        self.analysis_done.emit({
            "is_archive": False,
            "package_type": package_type,
            "engine_version": engine_version,
            "suggested_name": analysis.suggested_name,
            "resolved_asset_path": str(resolved_path),
            "plugin_folder_name": plugin_folder_name,
        })


class AddAssetAnalysisDialog(QDialog):
    """添加资产识别/分析详情窗口"""

    STAGES = ["分析资产版本", "分析资产类型"]

    def __init__(self, source_name: str, parent=None):
        super().__init__(parent)
        self._rows = []
        self._spinners = {}
        self.result = None

        self.setModal(True)
        self.setFixedWidth(420)
        self.setWindowTitle("识别分析")
        self.setWindowFlags(Qt.WindowType.Dialog | Qt.WindowType.FramelessWindowHint)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)

        container = QWidget(self)
        container.setObjectName("RenameResultContainer")
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.addWidget(container)

        layout = QVBoxLayout(container)
        layout.setContentsMargins(24, 20, 24, 18)
        layout.setSpacing(8)

        title = QLabel("资产识别分析中")
        title.setObjectName("RenameResultTitle")
        layout.addWidget(title)

        summary = QLabel(source_name)
        summary.setObjectName("RenameResultSummary")
        layout.addWidget(summary)

        self.detail_label = QLabel("等待开始...")
        self.detail_label.setObjectName("RenameResultTip")
        layout.addWidget(self.detail_label)

        layout.addSpacing(8)

        for idx, name in enumerate(self.STAGES):
            row = QHBoxLayout()
            row.setSpacing(10)

            icon = QLabel("○")
            icon.setObjectName("RenameStageIcon")
            icon.setFixedWidth(18)
            icon.setStyleSheet("color: #666666; font-size: 14px;")
            row.addWidget(icon)

            text = QLabel(f"阶段 {idx + 1}/{len(self.STAGES)} · {name}")
            text.setObjectName("RenameResultDoneItem")
            text.setStyleSheet("color: #666666;")
            row.addWidget(text, 1)

            layout.addLayout(row)
            self._rows.append((icon, text, row))

        layout.addSpacing(8)

    def set_stage_running(self, stage_idx: int, detail: str):
        if stage_idx < 0 or stage_idx >= len(self._rows):
            return
        icon, text, row = self._rows[stage_idx]
        spinner = _StageSpinnerLabel(self)
        spinner.setObjectName("RenameStageIcon")
        spinner.setFixedWidth(18)
        spinner.setStyleSheet("color: #4A9EFF; font-size: 14px;")

        row.removeWidget(icon)
        icon.hide()
        row.insertWidget(0, spinner)
        self._rows[stage_idx] = (spinner, text, row)
        self._spinners[stage_idx] = spinner

        text.setStyleSheet("color: #e0e0e0; font-weight: bold;")
        self.detail_label.setText(detail)

    def set_stage_done(self, stage_idx: int):
        if stage_idx < 0 or stage_idx >= len(self._rows):
            return
        icon, text, row = self._rows[stage_idx]
        spinner = self._spinners.pop(stage_idx, None)
        if spinner:
            spinner.stop()
            row.removeWidget(spinner)
            spinner.hide()

            done_icon = QLabel("✓")
            done_icon.setObjectName("RenameStageIcon")
            done_icon.setFixedWidth(18)
            done_icon.setStyleSheet("color: #4CAF50; font-size: 14px; font-weight: bold;")
            row.insertWidget(0, done_icon)
            self._rows[stage_idx] = (done_icon, text, row)

        text.setStyleSheet("color: #4CAF50;")

    def show_complete_then_close(self):
        """展示完成态，确保窗口至少显示 1.5 秒后再关闭"""
        self.detail_label.setText("识别分析完成")
        import time as _time
        elapsed_ms = int((_time.monotonic() - self._show_time) * 1000)
        remain_ms = max(0, 1500 - elapsed_ms)
        QTimer.singleShot(remain_ms, self.accept)

    def showEvent(self, event):
        import time as _time
        self._show_time = _time.monotonic()
        super().showEvent(event)
        self._center_on_content_area()

    def _center_on_content_area(self):
        """左右居中，上下偏上（距顶 1/4 处），相对右侧 RightPanel 区域"""
        self.adjustSize()
        host = self._find_right_panel()

        if host:
            host_global = host.mapToGlobal(host.rect().topLeft())
            x = host_global.x() + (host.width() - self.width()) // 2
            # 偏上：距顶部 1/4 处，让用户视觉舒适
            y = host_global.y() + (host.height() - self.height()) // 4
            self.move(x, y)
            return

        screen = QApplication.primaryScreen()
        if screen:
            available = screen.availableGeometry()
            x = available.x() + (available.width() - self.width()) // 2
            y = available.y() + (available.height() - self.height()) // 4
            self.move(x, y)

    def _find_right_panel(self):
        """向上遍历 parent 链，找到 objectName == 'RightPanel' 的组件（右侧整体区域）"""
        w = self.parent()
        while w is not None:
            if w.objectName() == "RightPanel":
                return w
            w = w.parent() if hasattr(w, 'parent') and callable(w.parent) else None
        return self.parent()


class AssetManagerUI(BaseModuleWidget):
    """资产管理器UI"""
    
    # 定义一个信号用于在主线程中重置按钮
    _reset_button_signal = pyqtSignal(str)  # 参数是资产名称
    # 定义信号用于更新预览按钮进度
    _update_preview_progress_signal = pyqtSignal(str, float)  # 资产名称, 进度(0.0-1.0)
    _update_preview_text_signal = pyqtSignal(str, str)  # 资产名称, 文本
    
    def __init__(self, logic, theme="dark", parent=None):
        super().__init__(parent)
        self.logic = logic
        self.theme = theme
        self.asset_cards = {}
        self.project_search_window = None  # 工程搜索窗口引用

        # 创建资产控制器，委托业务逻辑
        self.controller = AssetController(logic)

        # 强制使用紧凑视图
        self.current_view_mode = "compact"

        self.card_count = 0
        self._scroll_save_timer = QTimer()
        self._scroll_save_timer.setSingleShot(True)
        self._scroll_save_timer.timeout.connect(self._save_scroll_position)

        # ⚡ 性能优化：使用新的ThumbnailCache替代旧的LRU缓存
        self._thumbnail_cache = ThumbnailCache(max_size=300)
        self._thumbnail_cache.thumbnail_loaded.connect(self._on_thumbnail_loaded)
        self._thumbnail_cache.cache_stats_updated.connect(self._on_cache_stats_updated)

        self._init_ui()
        self._connect_signals()

        # 连接重置按钮信号
        self._reset_button_signal.connect(self._handle_reset_button)
        # 连接预览按钮更新信号
        self._update_preview_progress_signal.connect(self._handle_update_preview_progress)
        self._update_preview_text_signal.connect(self._handle_update_preview_text)

        # 初始化时主动加载主题样式（仅更新样式，不重新加载资产）
        self._apply_theme_styles(theme)

        # 不在初始化时加载资产，改为外部控制加载时机
        # 这样可以在启动界面显示期间异步加载，避免阻塞UI
        self._assets_loaded = False
        
        # 懒加载相关
        self._all_assets_data = []  # 存储所有资产数据
        self._loaded_card_count = 0  # 已加载的卡片数量
        self._initial_load_count = 50  # 初始加载数量
        self._load_more_count = 30  # 每次加载更多的数量
        self._is_loading_more = False  # 是否正在加载更多
        
        # 启用拖放功能
        self.setAcceptDrops(True)

    def showEvent(self, event):
        """首次显示时自动加载资产"""
        super().showEvent(event)
        if not self._assets_loaded:
            QTimer.singleShot(10, lambda: self.load_assets_async())
    
    def _init_ui(self):
        """初始化UI"""
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # 搜索和筛选区域
        filter_area = QWidget()
        filter_area.setObjectName("AssetFilterArea")
        filter_area.setMinimumHeight(60)  # 改为最小高度，允许下拉框展开
        filter_layout = QHBoxLayout(filter_area)
        filter_layout.setContentsMargins(20, 10, 20, 20)
        filter_layout.setSpacing(15)  # 增加下拉框之间的间距

        # 使用自定义delegate实现文本居中
        from PyQt6.QtWidgets import QStyledItemDelegate
        class CenterAlignDelegate(QStyledItemDelegate):
            def initStyleOption(self, option, index):
                super().initStyleOption(option, index)
                option.displayAlignment = Qt.AlignmentFlag.AlignCenter
        
        # 分类标签 + 选择框（使用 QHBoxLayout 控制间距）
        category_container = QWidget()
        category_layout = QHBoxLayout(category_container)
        category_layout.setContentsMargins(0, 0, 0, 0)
        category_layout.setSpacing(3)
        
        category_label = QLabel("分类：")
        category_label.setObjectName("AssetFilterLabel")
        category_label.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        category_layout.addWidget(category_label)
        
        self.category_filter = QComboBox()
        self.category_filter.setObjectName("AssetCategoryFilter")
        self.category_filter.setFixedHeight(36)
        self.category_filter.setMinimumWidth(120)
        self.category_filter.setMaximumWidth(180)
        self.category_filter.setFocusPolicy(Qt.FocusPolicy.ClickFocus)
        self.category_filter.setCursor(Qt.CursorShape.PointingHandCursor)
        self.category_filter.addItem("全部分类")  # 默认选项
        self.category_filter.currentTextChanged.connect(self._on_category_changed)
        self.category_filter.setItemDelegate(CenterAlignDelegate(self.category_filter))
        category_layout.addWidget(self.category_filter)
        
        filter_layout.addWidget(category_container)
        
        # 类型标签 + 筛选框（使用 QHBoxLayout 控制间距）
        type_container = QWidget()
        type_layout = QHBoxLayout(type_container)
        type_layout.setContentsMargins(0, 0, 0, 0)
        type_layout.setSpacing(3)
        
        type_label = QLabel("类型：")
        type_label.setObjectName("AssetFilterLabel")
        type_label.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        type_layout.addWidget(type_label)
        
        self.type_filter = QComboBox()
        self.type_filter.setObjectName("AssetTypeFilter")
        self.type_filter.setFixedHeight(36)
        self.type_filter.setMinimumWidth(100)
        self.type_filter.setMaximumWidth(140)
        self.type_filter.setFocusPolicy(Qt.FocusPolicy.ClickFocus)
        self.type_filter.setCursor(Qt.CursorShape.PointingHandCursor)
        self.type_filter.addItems(["全部类型", "Content 资产包", "UE 项目", "UE 插件", "其他资源"])
        self.type_filter.currentTextChanged.connect(self._on_type_filter_changed)
        self.type_filter.setItemDelegate(CenterAlignDelegate(self.type_filter))
        type_layout.addWidget(self.type_filter)
        
        filter_layout.addWidget(type_container)
        
        # 排序标签 + 选择框（使用 QHBoxLayout 控制间距）
        sort_container = QWidget()
        sort_layout = QHBoxLayout(sort_container)
        sort_layout.setContentsMargins(0, 0, 0, 0)
        sort_layout.setSpacing(3)
        
        sort_label = QLabel("排序：")
        sort_label.setObjectName("AssetFilterLabel")
        sort_label.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        sort_layout.addWidget(sort_label)
        
        self.sort_combo = QComboBox()
        self.sort_combo.setObjectName("AssetSortCombo")
        self.sort_combo.setFixedHeight(36)
        self.sort_combo.setMinimumWidth(120)
        self.sort_combo.setMaximumWidth(180)
        self.sort_combo.setFocusPolicy(Qt.FocusPolicy.ClickFocus)
        self.sort_combo.setCursor(Qt.CursorShape.PointingHandCursor)
        self.sort_combo.addItems([
            "🕒 最新添加",
            "🕒 最早添加",
            "🔤 A-Z",
            "🔤 Z-A"
        ])
        # 使用 activated 信号而不是 currentTextChanged，这样用户每次主动选择都会触发
        self.sort_combo.activated.connect(lambda: self._on_sort_changed(self.sort_combo.currentText()))
        self.sort_combo.setItemDelegate(CenterAlignDelegate(self.sort_combo))
        sort_layout.addWidget(self.sort_combo)
        
        filter_layout.addWidget(sort_container)

        # 添加弹性空间
        filter_layout.addStretch()
        
        # 搜索框
        self.search_input = QLineEdit()
        self.search_input.setObjectName("AssetSearchInput")
        self.search_input.setPlaceholderText("输入资产名称或拼音...")
        self.search_input.setFixedHeight(36)
        self.search_input.setMaximumWidth(200)
        self.search_input.setFocusPolicy(Qt.FocusPolicy.ClickFocus)  # 只有点击时才获得焦点
        self.search_input.textChanged.connect(self._on_search_changed)
        filter_layout.addWidget(self.search_input)
        
        # 添加资产按钮
        self.add_asset_btn = QPushButton("+ 添加资产")
        self.add_asset_btn.setObjectName("AddAssetButton")
        self.add_asset_btn.setFixedHeight(36)
        self.add_asset_btn.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        self.add_asset_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.add_asset_btn.setToolTip("添加资产到资产库")
        self.add_asset_btn.clicked.connect(self._show_add_asset_dialog)
        filter_layout.addWidget(self.add_asset_btn)

        # 移除视图切换按钮 - 统一使用紧凑视图

        main_layout.addWidget(filter_area)

        # 创建滚动区域
        scroll_area = QScrollArea()
        scroll_area.setObjectName("AssetScrollArea")
        scroll_area.setWidgetResizable(True)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll_area.setFocusPolicy(Qt.FocusPolicy.StrongFocus)  # 允许滚动区域接收焦点
        
        # 启用自动隐藏滚动条
        from core.utils.auto_hide_scrollbar import enable_auto_hide_scrollbar
        enable_auto_hide_scrollbar(scroll_area)

        # 滚动区域内容
        scroll_content = QWidget()
        scroll_content.setObjectName("AssetScrollContent")
        scroll_content.setFocusPolicy(Qt.FocusPolicy.StrongFocus)  # 允许内容区域接收焦点
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(0, 0, 0, 0)
        scroll_layout.setSpacing(20)

        # 创建网格容器（用于放置卡片）
        self.grid_widget = QWidget()
        self.grid_widget.setObjectName("AssetGridWidget")
        self.grid_layout = QGridLayout(self.grid_widget)
        self.grid_layout.setContentsMargins(0, 0, 0, 0)
        # 设置网格间距（紧凑视图）
        initial_spacing = 20
        self.grid_layout.setSpacing(initial_spacing)
        self.grid_layout.setAlignment(Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft)

        scroll_layout.addWidget(self.grid_widget)
        scroll_layout.addStretch()

        scroll_area.setWidget(scroll_content)
        main_layout.addWidget(scroll_area)

        # 保存滚动区域引用，用于后续事件处理
        self.scroll_area = scroll_area
        self.scroll_content = scroll_content
        
        # ⚡ 初始化时显示"加载中"占位符
        self._show_initial_loading_placeholder()

        # 安装事件过滤器，用于处理点击事件
        scroll_area.viewport().installEventFilter(self)
        scroll_content.installEventFilter(self)
        
        # 连接滚动条信号，实现懒加载 + 保存滚动位置
        scroll_area.verticalScrollBar().valueChanged.connect(self._on_scroll_changed)
        scroll_area.verticalScrollBar().valueChanged.connect(self._on_scroll_position_changed)

    def eventFilter(self, obj, event):
        """事件过滤器：处理点击事件以清除搜索框焦点"""
        from PyQt6.QtCore import QEvent
        from PyQt6.QtGui import QMouseEvent

        # 如果是鼠标按下事件，并且点击的不是搜索框
        if event.type() == QEvent.Type.MouseButtonPress:
            if isinstance(event, QMouseEvent):
                # 清除搜索框焦点
                if self.search_input.hasFocus():
                    self.search_input.clearFocus()

        return super().eventFilter(obj, event)

    def _connect_signals(self):
        """连接信号"""
        if self.logic:
            self.logic.assets_loaded.connect(self._on_assets_loaded)
            self.logic.error_occurred.connect(self._on_error_occurred)
            self.logic.thumbnail_updated.connect(self._on_thumbnail_updated)

    def _load_categories(self):
        """加载分类列表到下拉框"""
        try:
            # 通过控制器获取所有分类
            categories = self.controller.get_categories()
            
            # 保存当前选中的分类
            current_category = self.category_filter.currentText()
            
            # 阻塞信号，避免在重新填充时触发保存
            self.category_filter.blockSignals(True)
            
            # 清空并重新填充（保留"全部分类"）
            self.category_filter.clear()
            self.category_filter.addItem("全部分类")
            
            # 添加所有分类
            for category in categories:
                self.category_filter.addItem(category)
            
            # 添加"⚙️ 分类管理"选项
            self.category_filter.addItem("⚙️ 分类管理")
            
            # 优先恢复已保存的分类，否则恢复之前的选择
            saved_category = self.controller.load_ui_state("selected_category", "")
            restore_target = saved_category if saved_category else current_category
            if restore_target and restore_target != "⚙️ 分类管理":
                index = self.category_filter.findText(restore_target)
                if index >= 0:
                    self.category_filter.setCurrentIndex(index)
                    # 同步到控制器
                    self.controller.set_category(restore_target)
            
            # 恢复信号
            self.category_filter.blockSignals(False)
            
            logger.info(f"已加载 {len(categories)} 个分类")
        except Exception as e:
            logger.error(f"加载分类失败: {e}", exc_info=True)
            self.category_filter.blockSignals(False)
    

    
    def _on_category_changed(self, category: str):
        """分类选择改变事件"""
        logger.info(f"分类选择改变: '{category}'")
        
        # 如果选中"⚙️ 分类管理"，显示分类管理对话框
        if category == "⚙️ 分类管理":
            self._show_add_category_dialog()
            # 恢复到"全部分类"
            self.category_filter.setCurrentIndex(0)
            return
        
        # 通过控制器更新分类过滤
        self.controller.set_category(category)
        self.controller.save_ui_state("selected_category", category)
        self._apply_filter_to_ui()
    
    def _on_search_changed(self, search_text: str):
        """搜索文本改变事件（实时搜索，带防抖优化）"""
        self.controller.set_search_text(search_text)
        
        # 添加防抖，避免用户快速输入时频繁过滤
        if not hasattr(self, '_search_timer'):
            self._search_timer = QTimer()
            self._search_timer.setSingleShot(True)
            self._search_timer.timeout.connect(self._perform_search)
        
        # 停止之前的定时器
        self._search_timer.stop()
        
        # 如果搜索文本为空，立即执行（显示全部）
        if not self.controller.search_text:
            self._perform_search()
        else:
            # 否则延迟300ms执行，等待用户输入完成
            self._search_timer.start(300)
    
    def _perform_search(self):
        """执行实际的搜索过滤"""
        logger.info(f"执行搜索: '{self.controller.search_text}'")
        self._apply_filter_to_ui()
    
    def _on_type_filter_changed(self, type_name: str):
        """类型筛选改变事件"""
        # 中文到英文类型映射
        type_map = {
            "全部类型": "全部类型",
            "Content 资产包": "Content",
            "UE 项目": "Project",
            "UE 插件": "Plugin",
            "其他资源": "Others"
        }
        # 转换为英文类型名
        english_type = type_map.get(type_name, type_name)
        self.controller.set_type_filter(english_type)
        self.controller.save_ui_state("selected_type", type_name)
        logger.info(f"类型筛选改变: {type_name} -> {english_type}")
        self._apply_filter_to_ui()
    


    def _on_sort_changed(self, sort_method: str):
        """排序方式改变事件"""
        self.controller.set_sort_method(sort_method)
        self.controller.save_ui_state("sort_method", sort_method)
        logger.info(f"排序方式改变: {sort_method}")
        
        # 重新应用筛选和排序
        self._apply_filter_to_ui()

    def _apply_filter_to_ui(self):
        """通过控制器获取过滤结果，并更新 UI 卡片的可见性和布局"""
        try:
            # 通过控制器获取过滤后的资产列表
            matched_assets = self.controller.get_filtered_assets()
            
            # 检查是否为空状态
            if len(matched_assets) == 0:
                # 判断是真的为空还是搜索/筛选结果为空
                total_assets = len(self.logic.assets) if self.logic else 0
                
                if total_assets == 0:
                    # 资产库真的为空
                    logger.info("资产库为空，显示空状态占位符")
                    # 隐藏所有卡片
                    for card in self.asset_cards.values():
                        card.setVisible(False)
                    # 隐藏网格容器
                    self.grid_widget.hide()
                    # 显示空状态占位符
                    self._show_empty_state_placeholder()
                else:
                    # 搜索/筛选结果为空
                    logger.info("搜索/筛选结果为空，显示无结果提示")
                    # 隐藏所有卡片
                    for card in self.asset_cards.values():
                        card.setVisible(False)
                    # 隐藏网格容器
                    self.grid_widget.hide()
                    # 显示无结果占位符
                    self._show_no_results_placeholder()
                return
            
            # 移除空状态占位符（如果存在）
            if hasattr(self, '_empty_state_container') and self._empty_state_container:
                try:
                    self._empty_state_container.deleteLater()
                    self._empty_state_container = None
                except:
                    pass
            
            # 移除无结果占位符（如果存在）
            if hasattr(self, '_no_results_container') and self._no_results_container:
                try:
                    self._no_results_container.deleteLater()
                    self._no_results_container = None
                except:
                    pass
            
            # 显示网格容器
            self.grid_widget.show()
            
            # 先隐藏所有卡片
            for card in self.asset_cards.values():
                card.setVisible(False)
            
            # 按顺序重新排列匹配的卡片
            # 懒加载优化：只显示前N个，其余的在滚动时加载
            visible_count = 0
            max_initial_visible = self._initial_load_count  # 初始只显示这么多
            
            for asset in matched_assets:
                card = self.asset_cards.get(asset.id)
                if card:
                    # 只有前N个才立即显示和重新布局
                    if visible_count < max_initial_visible:
                        # 计算新的网格位置（紧凑视图：5列）
                        row = visible_count // 5
                        col = visible_count % 5
                        
                        # 移除旧位置
                        self.grid_layout.removeWidget(card)
                        # 添加到新位置
                        self.grid_layout.addWidget(card, row, col)
                        # 显示卡片
                        card.setVisible(True)
                    
                    visible_count += 1
            
            logger.info(f"搜索结果: 显示 {visible_count} 个资产（初始显示 {min(visible_count, max_initial_visible)} 个）")
            
            # 触发可见缩略图的加载
            QTimer.singleShot(100, self._load_visible_thumbnails)
            
        except Exception as e:
            logger.error(f"过滤资产失败: {e}", exc_info=True)

    def _refresh_assets(self):
        """刷新资产列表（应用搜索过滤，使用批量优化 + 延迟加载缩略图）"""
        if not self.logic:
            logger.warning("Logic 层未初始化")
            return

        try:
            # 通过控制器获取过滤后的资产列表
            assets = self.controller.get_filtered_assets()
            logger.info(f"刷新资产列表: {len(assets)} 个资产")

            # ⚡ 性能优化：禁用更新，避免每次添加卡片时都重绘
            self.grid_widget.setUpdatesEnabled(False)

            try:
                # 清空现有卡片
                self._clear_cards()
                self.card_count = 0

                # 收集需要延迟加载缩略图的卡片
                cards_to_load = []

                # 通过控制器转换 Asset 对象为字典格式并批量创建卡片
                for i, asset in enumerate(assets):
                    asset_dict = self.controller.convert_asset_to_dict(asset)
                    card = self._add_asset_card(asset_dict, i, defer_thumbnail=True)
                    if card:
                        cards_to_load.append(card)
            finally:
                # ⚡ 重新启用更新，一次性重绘所有卡片
                self.grid_widget.setUpdatesEnabled(True)
                logger.debug(f"批量刷新了 {len(assets)} 个资产卡片")

            # ⚡ 延迟加载缩略图，避免阻塞 UI
            if cards_to_load:
                QTimer.singleShot(0, self._load_visible_thumbnails)

        except Exception as e:
            logger.error(f"刷新资产列表失败: {e}", exc_info=True)

    def _apply_theme_styles(self, theme: str):
        """应用主题样式（不重新加载资产）
        
        Args:
            theme: 主题名称 ("dark" 或 "light")
        """
        self.theme = theme
        
        # 不需要手动应用主题，依赖应用程序级别的QSS
        # 因为主窗口已经通过 style_system.apply_theme(app, theme_name) 应用了全局主题
        # 如果我们再调用 widget.setStyleSheet()，会覆盖全局样式
        
        logger.info(f"资产管理器主题样式已更新为: {theme}")
    
    def update_theme(self, theme: str):
        """更新主题

        Args:
            theme: 主题名称 ("dark" 或 "light")
        """
        # 应用主题样式
        self._apply_theme_styles(theme)

        # 只更新现有卡片的主题，不重新创建卡片（避免卡顿）
        for card in self.asset_cards.values():
            if hasattr(card, 'update_theme'):
                card.update_theme(theme)

        logger.info(f"资产管理器主题已更新为: {theme}")
    
    def on_theme_changed(self, theme_name: str) -> None:
        """主题切换回调方法（继承自 BaseModuleWidget）
        
        当应用主题切换时，此方法会被主窗口调用。
        
        Args:
            theme_name: 新主题名称 ('dark' 或 'light')
        """
        # 调用基类实现以刷新样式
        super().on_theme_changed(theme_name)
        
        # 调用现有的 update_theme 方法更新资产卡片
        self.update_theme(theme_name)
        
        logger.debug(f"AssetManagerUI 主题切换完成: {theme_name}")
    
    def load_assets_async(self, on_complete=None, force_reload=False):
        """异步加载资产数据
        
        Args:
            on_complete: 加载完成回调函数
            force_reload: 是否强制重新加载（忽略已加载标志）
        """
        logger.info(f"开始异步加载资产，已加载状态: {self._assets_loaded}, 强制重载: {force_reload}")
        
        if self._assets_loaded and not force_reload:
            logger.info("资产已加载，直接返回")
            if on_complete:
                on_complete()
            return
        
        # ⚡ 显示"加载中"提示
        self._show_loading_placeholder()
            
        # 使用QTimer延迟执行，避免阻塞UI
        from PyQt6.QtCore import QTimer
        QTimer.singleShot(10, lambda: self._load_assets_async(on_complete))
    
    def _create_loading_widget(self):
        """创建统一的加载中占位符控件"""
        loading_container = QWidget()
        loading_container.setObjectName("assetManagerEmptyContainer")
        loading_layout = QVBoxLayout()
        loading_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # 使用 QProgressBar 不确定模式做旋转动画
        from PyQt6.QtWidgets import QProgressBar
        spinner = QProgressBar()
        spinner.setRange(0, 0)  # 不确定模式
        spinner.setTextVisible(False)
        spinner.setFixedSize(200, 3)
        spinner.setStyleSheet("""
            QProgressBar {
                border: none;
                border-radius: 1px;
                background-color: rgba(255, 255, 255, 0.05);
            }
            QProgressBar::chunk {
                background-color: rgba(255, 255, 255, 0.2);
                border-radius: 1px;
            }
        """)
        loading_layout.addWidget(spinner, 0, Qt.AlignmentFlag.AlignCenter)
        
        loading_layout.addSpacing(12)
        
        loading_label = QLabel("正在加载资产...")
        loading_label.setObjectName("assetManagerEmptyLabel")
        loading_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        loading_layout.addWidget(loading_label)
        
        loading_container.setLayout(loading_layout)
        return loading_container
    
    def _show_loading_placeholder(self):
        """显示加载中占位符"""
        try:
            self.grid_widget.hide()
            
            if hasattr(self, '_loading_container') and self._loading_container:
                try:
                    self._loading_container.deleteLater()
                    self._loading_container = None
                except:
                    pass
            
            loading_container = self._create_loading_widget()
            self.scroll_content.layout().insertWidget(0, loading_container)
            self._loading_container = loading_container
        except Exception as e:
            logger.error(f"显示加载占位符失败: {e}")
    
    def _show_initial_loading_placeholder(self):
        """初始化时显示加载中占位符"""
        try:
            self.grid_widget.hide()
            
            loading_container = self._create_loading_widget()
            self.scroll_content.layout().insertWidget(0, loading_container)
            self._loading_container = loading_container
        except Exception as e:
            logger.error(f"显示初始加载占位符失败: {e}")
    
    def _create_empty_state_widget(self):
        """创建空状态占位符控件"""
        empty_container = QWidget()
        empty_container.setObjectName("assetManagerEmptyContainer")
        empty_layout = QVBoxLayout()
        empty_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        empty_layout.setSpacing(20)
        
        # 图标 📦
        icon_label = QLabel("📦")
        icon_label.setObjectName("assetManagerEmptyIcon")
        icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        empty_layout.addWidget(icon_label)
        
        # 主标题
        title_label = QLabel("资产库为空")
        title_label.setObjectName("assetManagerEmptyLabel")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        empty_layout.addWidget(title_label)
        
        empty_layout.addSpacing(10)
        
        # 功能亮点列表
        features_layout = QVBoxLayout()
        features_layout.setSpacing(8)
        features_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        features = [
            "✨ 智能识别资产类型（内容包、插件、工程、模型、图片等）",
            "📦 支持压缩包直接添加，自动解压和包装",
            "🎯 拖入即可，无需手动解压"
        ]
        
        for feature in features:
            feature_label = QLabel(feature)
            feature_label.setObjectName("assetManagerEmptyHint")
            feature_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            features_layout.addWidget(feature_label)
        
        empty_layout.addLayout(features_layout)
        
        empty_layout.addSpacing(10)
        
        # 操作引导
        hint_label = QLabel("拖入资产文件或点击添加按钮开始使用")
        hint_label.setObjectName("assetManagerEmptyHint")
        hint_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        empty_layout.addWidget(hint_label)
        
        empty_container.setLayout(empty_layout)
        return empty_container
    
    def _create_no_results_widget(self):
        """创建搜索/筛选无结果占位符控件"""
        no_results_container = QWidget()
        no_results_container.setObjectName("assetManagerEmptyContainer")
        no_results_layout = QVBoxLayout()
        no_results_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        no_results_layout.setSpacing(20)
        
        # 图标 🔍
        icon_label = QLabel("🔍")
        icon_label.setObjectName("assetManagerEmptyIcon")
        icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        no_results_layout.addWidget(icon_label)
        
        # 主标题
        title_label = QLabel("未找到匹配的资产")
        title_label.setObjectName("assetManagerEmptyLabel")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        no_results_layout.addWidget(title_label)
        
        no_results_layout.addSpacing(10)
        
        # 提示信息
        hint_label = QLabel("尝试调整搜索关键词或筛选条件")
        hint_label.setObjectName("assetManagerEmptyHint")
        hint_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        no_results_layout.addWidget(hint_label)
        
        no_results_container.setLayout(no_results_layout)
        return no_results_container
    
    def _show_empty_state_placeholder(self):
        """显示空状态占位符"""
        try:
            self.grid_widget.hide()
            
            # 移除旧的占位符
            if hasattr(self, '_empty_state_container') and self._empty_state_container:
                try:
                    self._empty_state_container.deleteLater()
                    self._empty_state_container = None
                except:
                    pass
            
            # 创建并显示空状态占位符
            empty_container = self._create_empty_state_widget()
            self.scroll_content.layout().insertWidget(0, empty_container)
            self._empty_state_container = empty_container
            
            logger.info("显示空状态占位符")
        except Exception as e:
            logger.error(f"显示空状态占位符失败: {e}")
    
    def _show_no_results_placeholder(self):
        """显示搜索/筛选无结果占位符"""
        try:
            self.grid_widget.hide()
            
            # 移除旧的占位符
            if hasattr(self, '_no_results_container') and self._no_results_container:
                try:
                    self._no_results_container.deleteLater()
                    self._no_results_container = None
                except:
                    pass
            
            # 创建并显示无结果占位符
            no_results_container = self._create_no_results_widget()
            self.scroll_content.layout().insertWidget(0, no_results_container)
            self._no_results_container = no_results_container
            
            logger.info("显示搜索/筛选无结果占位符")
        except Exception as e:
            logger.error(f"显示无结果占位符失败: {e}")
    def _create_empty_state_widget(self):
        """创建空状态占位符控件"""
        empty_container = QWidget()
        empty_container.setObjectName("assetManagerEmptyContainer")
        empty_layout = QVBoxLayout()
        empty_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        empty_layout.setSpacing(20)

        # 图标 📦
        icon_label = QLabel("📦")
        icon_label.setObjectName("assetManagerEmptyIcon")
        icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        empty_layout.addWidget(icon_label)

        # 主标题
        title_label = QLabel("资产库为空")
        title_label.setObjectName("assetManagerEmptyLabel")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        empty_layout.addWidget(title_label)

        empty_layout.addSpacing(10)

        # 功能亮点列表
        features_layout = QVBoxLayout()
        features_layout.setSpacing(8)
        features_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        features = [
            "✨ 智能识别资产类型（内容包、插件、工程、模型、图片等）",
            "📦 支持压缩包直接添加，自动解压和包装",
            "🎯 拖入即可，无需手动解压"
        ]

        for feature in features:
            feature_label = QLabel(feature)
            feature_label.setObjectName("assetManagerEmptyHint")
            feature_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            features_layout.addWidget(feature_label)

        empty_layout.addLayout(features_layout)

        empty_layout.addSpacing(10)

        # 操作引导
        hint_label = QLabel("拖入资产文件或点击添加按钮开始使用")
        hint_label.setObjectName("assetManagerEmptyHint")
        hint_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        empty_layout.addWidget(hint_label)

        empty_container.setLayout(empty_layout)
        return empty_container

    def _show_empty_state_placeholder(self):
        """显示空状态占位符"""
        try:
            self.grid_widget.hide()

            # 移除旧的占位符
            if hasattr(self, '_empty_state_container') and self._empty_state_container:
                try:
                    self._empty_state_container.deleteLater()
                    self._empty_state_container = None
                except:
                    pass

            # 创建并显示空状态占位符
            empty_container = self._create_empty_state_widget()
            self.scroll_content.layout().insertWidget(0, empty_container)
            self._empty_state_container = empty_container

            logger.info("显示空状态占位符")
        except Exception as e:
            logger.error(f"显示空状态占位符失败: {e}")

    
    def _load_assets_async(self, on_complete=None):
        """实际执行异步加载资产"""
        logger.info("开始执行异步加载资产")
        
        # ⚡ 移除初始加载占位符
        if hasattr(self, '_loading_container') and self._loading_container:
            try:
                self._loading_container.deleteLater()
                self._loading_container = None
            except:
                pass
        
        # 先隐藏网格容器，避免创建过程中卡片闪现
        self.grid_widget.hide()

        # 通过 logic 层获取资产列表
        if not self.logic:
            logger.warning("Logic 层未初始化")
            if on_complete:
                on_complete()
            return

        try:
            # 从 logic 层获取资产列表
            assets = self.logic.assets
            logger.info(f"从 logic 层加载了 {len(assets)} 个资产")

            # 清空现有卡片（使用优化的清空方法）
            self._clear_cards()
            self.card_count = 0
            self._loaded_card_count = 0

            # 立即强制处理事件，让清空操作立即生效
            from PyQt6.QtWidgets import QApplication
            QApplication.processEvents()

            # 转换 Asset 对象为字典格式（通过控制器）
            assets_data = []
            for asset in assets:
                asset_dict = self.controller.convert_asset_to_dict(asset)
                assets_data.append(asset_dict)

            # 保存所有资产数据
            self._all_assets_data = assets_data
            
            # 只加载初始数量的卡片
            initial_count = min(self._initial_load_count, len(assets_data))
            logger.info(f"懒加载：初始加载 {initial_count}/{len(assets_data)} 个资产卡片")
            
            # 立即开始创建第一批卡片
            self._create_assets_batch(assets_data[:initial_count], 0, on_complete)
            
            # 更新已加载计数
            self._loaded_card_count = initial_count

        except Exception as e:
            logger.error(f"加载资产失败: {e}", exc_info=True)
            if on_complete:
                on_complete()
    
    def _create_assets_batch(self, assets, start_index, on_complete):
        """分批创建资产卡片"""
        # 检查是否为空状态
        if len(assets) == 0:
            logger.info("资产库为空，显示空状态占位符")
            self._show_empty_state_placeholder()
            self._assets_loaded = True
            if on_complete:
                on_complete()
            return
        
        batch_size = 30  # 进一步增加批次大小，减少批次切换开销
        end_index = min(start_index + batch_size, len(assets))

        logger.info(f"创建资产卡片批次: {start_index}-{end_index-1}/{len(assets)}")

        # 第一批时禁用网格更新，避免卡片在创建过程中闪现
        if start_index == 0:
            self.grid_widget.setUpdatesEnabled(False)
            
            # 移除空状态占位符（如果存在）
            if hasattr(self, '_empty_state_container') and self._empty_state_container:
                try:
                    self._empty_state_container.deleteLater()
                    self._empty_state_container = None
                except:
                    pass

        # 创建当前批次的卡片
        for i in range(start_index, end_index):
            self._add_asset_card(assets[i], i)

        # 如果还有更多资产需要创建，继续创建下一批
        if end_index < len(assets):
            # 使用0ms间隔，立即创建下一批，但通过QTimer让出控制权给UI线程
            from PyQt6.QtCore import QTimer
            QTimer.singleShot(0, lambda: self._create_assets_batch(assets, end_index, on_complete))
        else:
            # 所有资产创建完成
            logger.info("所有资产卡片创建完成")
            self._assets_loaded = True
            
            # 重新启用网格更新并显示
            self.grid_widget.setUpdatesEnabled(True)
            self.grid_widget.show()
            
            # 加载分类列表（内部会恢复已保存分类）
            self._load_categories()
            
            # 恢复已保存的类型筛选
            saved_type = self.controller.load_ui_state("selected_type", "")
            if saved_type:
                idx = self.type_filter.findText(saved_type)
                if idx >= 0:
                    self.type_filter.setCurrentIndex(idx)

            # 恢复已保存的排序方式
            saved_sort = self.controller.load_ui_state("sort_method", "")
            if saved_sort:
                idx = self.sort_combo.findText(saved_sort)
                if idx >= 0:
                    self.sort_combo.blockSignals(True)
                    self.sort_combo.setCurrentIndex(idx)
                    self.sort_combo.blockSignals(False)
                    self.controller.set_sort_method(saved_sort)
            
            # 应用筛选和排序（无论是否有保存的排序方式都要调用）
            self._apply_filter_to_ui()

            # 恢复滚动位置（延迟等待布局完成）
            from PyQt6.QtCore import QTimer as QtTimer
            saved_scroll = self.controller.load_ui_state("scroll_position", 0)
            if saved_scroll:
                QtTimer.singleShot(300, lambda: self.scroll_area.verticalScrollBar().setValue(saved_scroll))
            
            # 触发初始可见缩略图的加载
            QtTimer.singleShot(100, self._load_visible_thumbnails)

            # 触发后台增量扫描，检测缓存外的变化
            if self.logic and hasattr(self.logic, 'rescan_in_background'):
                if not getattr(self, '_background_rescanning', False):
                    QtTimer.singleShot(500, self._trigger_background_rescan)
            
            if on_complete:
                on_complete()
    
    def _trigger_background_rescan(self):
        """触发后台增量扫描，扫描完成后刷新 UI"""
        if not self.logic:
            return
        self._background_rescanning = True
        self.logic.rescan_in_background()

        # 延迟检查扫描结果（扫描完成后 logic.assets 会更新）
        from PyQt6.QtCore import QTimer
        self._rescan_check_timer = QTimer()
        self._rescan_check_timer.setSingleShot(True)
        self._rescan_check_timer.timeout.connect(self._check_rescan_result)
        self._rescan_check_timer.start(5000)  # 5 秒后检查

    def _check_rescan_result(self):
        """检查后台扫描是否导致资产数量变化，如有变化则刷新 UI"""
        self._background_rescanning = False
        if not self.logic:
            return
        current_count = len(self._all_assets_data)
        new_count = len(self.logic.assets)
        if current_count != new_count:
            logger.info(f"增量扫描检测到变化: {current_count} -> {new_count}，刷新 UI")
            self._assets_loaded = False
            self.load_assets_async(force_reload=True)

    def _on_scroll_changed(self, value):
        """滚动条位置改变时触发 - 实现懒加载和缩略图按需加载"""
        # 1. 懒加载更多资产卡片
        if not self._is_loading_more and self._loaded_card_count < len(self._all_assets_data):
            scrollbar = self.scroll_area.verticalScrollBar()
            max_value = scrollbar.maximum()
            
            # 当滚动到距离底部200px时，加载更多
            if max_value - value < 200:
                self._load_more_assets()
        
        # 2. 按需加载可见卡片的缩略图（防抖）
        if hasattr(self, '_thumbnail_load_timer'):
            self._thumbnail_load_timer.stop()
        else:
            self._thumbnail_load_timer = QTimer()
            self._thumbnail_load_timer.setSingleShot(True)
            self._thumbnail_load_timer.timeout.connect(self._load_visible_thumbnails)
        
        # 延迟100ms加载，避免滚动时频繁触发
        self._thumbnail_load_timer.start(100)
    
    def _load_more_assets(self):
        """加载更多资产卡片"""
        if self._is_loading_more:
            return
        
        # 计算要加载的资产范围
        start_index = self._loaded_card_count
        end_index = min(start_index + self._load_more_count, len(self._all_assets_data))
        
        if start_index >= end_index:
            return
        
        self._is_loading_more = True
        logger.info(f"懒加载：加载更多资产 {start_index}-{end_index}/{len(self._all_assets_data)}")
        
        # 获取要加载的资产数据
        assets_to_load = self._all_assets_data[start_index:end_index]
        
        # 创建卡片
        for i, asset_data in enumerate(assets_to_load):
            self._add_asset_card(asset_data, start_index + i)
        
        # 更新已加载计数
        self._loaded_card_count = end_index
        self._is_loading_more = False
        
        logger.info(f"懒加载：已加载 {self._loaded_card_count}/{len(self._all_assets_data)} 个资产")
        
        # 加载新创建卡片的缩略图（如果可见）
        QTimer.singleShot(0, self._load_visible_thumbnails)
    
    def _load_visible_thumbnails(self):
        """加载可见区域的缩略图（使用ThumbnailCache）"""
        if not hasattr(self, 'scroll_area') or not self.scroll_area:
            return
        
        # 获取滚动区域的可见矩形
        viewport = self.scroll_area.viewport()
        viewport_rect = viewport.rect()
        
        # 获取滚动偏移
        scroll_offset = self.scroll_area.verticalScrollBar().value()
        
        # 遍历所有卡片，检查可见性
        visible_cards = []
        for asset_id, card in self.asset_cards.items():
            if not card or not card.isVisible():
                continue
            
            # 获取卡片在滚动内容中的位置
            card_pos = card.pos()
            card_rect = card.rect()
            
            # 计算卡片相对于视口的位置
            card_top = card_pos.y() - scroll_offset
            card_bottom = card_top + card_rect.height()
            
            # 检查卡片是否在可见区域内（包含一些预加载边距）
            preload_margin = 300  # 预加载边距（像素）
            if card_bottom >= -preload_margin and card_top <= viewport_rect.height() + preload_margin:
                visible_cards.append((asset_id, card))
        
        # 使用ThumbnailCache加载可见卡片的缩略图
        if visible_cards:
            logger.debug(f"检测到 {len(visible_cards)} 个可见卡片，开始加载缩略图")
            for asset_id, card in visible_cards:
                # 检查卡片是否已经加载过缩略图
                if hasattr(card, '_thumbnail_loaded') and card._thumbnail_loaded:
                    continue
                
                # 获取缩略图路径
                thumbnail_path = None
                if hasattr(card, 'thumbnail_path') and card.thumbnail_path:
                    thumbnail_path = Path(card.thumbnail_path)
                
                # 只有当缩略图路径存在时才使用缓存加载
                if thumbnail_path and thumbnail_path.exists():
                    # 从缓存获取缩略图（如果不在缓存中会触发异步加载）
                    pixmap = self._thumbnail_cache.get(asset_id, thumbnail_path)
                    
                    # 如果返回的不是占位符（即缓存命中），立即更新卡片
                    if pixmap and not pixmap.isNull():
                        self._update_card_thumbnail(card, pixmap)
                else:
                    # 没有缩略图路径，显示默认文字
                    if hasattr(card, '_show_default_icon'):
                        card._show_default_icon()
                        card._thumbnail_loaded = True  # 标记为已加载，避免重复处理
    
    def _on_thumbnail_loaded(self, asset_id: str, pixmap: QPixmap):
        """缩略图加载完成回调
        
        Args:
            asset_id: 资产ID
            pixmap: 加载的缩略图
        """
        # 查找对应的卡片并更新缩略图
        if asset_id in self.asset_cards:
            card = self.asset_cards[asset_id]
            self._update_card_thumbnail(card, pixmap)
            logger.debug(f"更新卡片缩略图: {asset_id}")
    
    def _update_card_thumbnail(self, card, pixmap: QPixmap):
        """更新卡片的缩略图
        
        Args:
            card: 资产卡片对象
            pixmap: 缩略图QPixmap
        """
        if not card or pixmap.isNull():
            return
        
        try:
            # 根据卡片类型缩放缩略图
            if isinstance(card, ModernAssetCard):
                target_w, target_h = 212, 153
            else:  # CompactAssetCard
                target_w, target_h = 172, 110
            
            # 缩放缩略图
            scaled_pixmap = pixmap.scaled(
                target_w, target_h,
                Qt.AspectRatioMode.KeepAspectRatioByExpanding,
                Qt.TransformationMode.SmoothTransformation
            )
            
            # 更新卡片的缩略图
            if hasattr(card, 'thumbnail_label') and card.thumbnail_label:
                card.thumbnail_label.setPixmap(scaled_pixmap)
                card._thumbnail_loaded = True
        except Exception as e:
            logger.error(f"更新卡片缩略图失败: {e}", exc_info=True)
    
    def _on_cache_stats_updated(self, size: int, hits: int, misses: int):
        """缓存统计信息更新回调
        
        Args:
            size: 当前缓存大小
            hits: 缓存命中次数
            misses: 缓存未命中次数
        """
        # 可以在这里添加统计信息的显示或日志
        total = hits + misses
        hit_rate = (hits / total * 100) if total > 0 else 0
        logger.debug(f"缩略图缓存统计: 大小={size}, 命中率={hit_rate:.1f}% ({hits}/{total})")

    
    def _load_assets(self):
        """同步加载资产数据（保留用于主题切换等场景）"""
        if not self.logic:
            logger.warning("Logic 层未初始化")
            return

        try:
            assets = self.logic.assets
            logger.info(f"从 logic 层加载了 {len(assets)} 个资产")

            # 清空现有卡片
            self._clear_cards()
            self.card_count = 0

            # 通过控制器转换 Asset 对象为字典格式
            for i, asset in enumerate(assets):
                asset_dict = self.controller.convert_asset_to_dict(asset)
                self._add_asset_card(asset_dict, i)

        except Exception as e:
            logger.error(f"加载资产失败: {e}", exc_info=True)
    
    def _add_asset_card(self, asset_data, index, defer_thumbnail=True):
        """添加资产卡片

        Args:
            asset_data: 资产数据字典
            index: 索引
            defer_thumbnail: 是否延迟加载缩略图（默认True，使用懒加载）

        Returns:
            创建的卡片对象（用于延迟加载缩略图）
        """
        name = asset_data.get('name', '未命名资产')
        category = asset_data.get('category', '默认分类')
        size = format_size(asset_data.get('size', 0))
        thumbnail_path = asset_data.get('thumbnail_path', '')
        asset_type = "资源包" if asset_data.get('asset_type') == 'package' else "文件"
        created_time = format_time(asset_data.get('created_time', ''))
        asset_path = asset_data.get('path', '')  # 获取资产路径
        engine_min_version = asset_data.get('engine_min_version', '')  # 获取引擎版本
        package_type = asset_data.get('package_type', 'content')  # 获取包装类型

        # 检查是否有文档（通过控制器）
        asset_id = asset_data.get('id', '')
        has_document = self.controller.check_asset_has_document(asset_id) if asset_id else False

        # 根据当前视图模式创建卡片（强制使用紧凑视图）
        card = CompactAssetCard(name, category, thumbnail_path, asset_type,
                               theme=self.theme, defer_thumbnail=defer_thumbnail,
                               asset_path=asset_path, engine_min_version=engine_min_version,
                               package_type=package_type, asset_size=size,
                               parent=self)
        row = index // 5
        col = index % 5

        # 设置卡片的asset_id和package_type
        if asset_id:
            card.asset_id = asset_id
        
        # 存储 package_type 用于预览行为分化
        card.package_type = asset_data.get('package_type', 'content')
        
        # 根据 package_type 调整预览按钮
        if hasattr(card, 'preview_btn'):
            if card.package_type == 'plugin':
                card.preview_btn.update_button_text("不可预览")
                card.preview_btn.setEnabled(False)
                card.preview_btn.setCursor(Qt.CursorShape.ForbiddenCursor)
            elif card.package_type == 'project':
                card.preview_btn.update_button_text("▶  打开项目")
        
        # 连接信号（使用Qt.ConnectionType.QueuedConnection确保信号正确传递）
        card.preview_clicked.connect(lambda: self._on_preview_asset(name))
        card.edit_info_requested.connect(self._on_edit_asset_info)
        # open_path_requested 信号不再需要，直接在卡片内部处理
        card.delete_requested.connect(self._on_delete_asset)
        card.detail_requested.connect(self._on_detail_requested, Qt.ConnectionType.QueuedConnection)  # 连接详情请求信号
        card.import_requested.connect(self._on_import_asset)  # 连接导入请求信号
        card.export_requested.connect(self._on_export_asset)  # 连接导出请求信号

        # 添加到网格布局
        self.grid_layout.addWidget(card, row, col)
        self.card_count += 1
        
        # 将卡片添加到字典（用于快速搜索和隐藏/显示）
        if asset_id:
            self.asset_cards[asset_id] = card

        return card

    def _load_view_mode(self) -> str:
        """从配置文件加载视图模式（委托给控制器）

        Returns:
            str: 视图模式 ("detailed" 或 "compact")
        """
        return self.controller.load_view_mode()

    def _save_view_mode(self):
        """保存视图模式到配置文件（委托给控制器）"""
        self.controller.save_view_mode(self.current_view_mode)

    def _update_view_toggle_button(self):
        """更新视图切换按钮的状态"""
        if self.current_view_mode == "detailed":
            self.view_toggle_btn.setText("≡")
            self.view_toggle_btn.setToolTip("切换到简约视图")
        else:
            self.view_toggle_btn.setText("⊞")
            self.view_toggle_btn.setToolTip("切换到详细视图")

    def _toggle_view_mode(self):
        """切换视图模式"""
        # 防抖：避免快速重复点击
        import time
        current_time = time.time() * 1000
        if hasattr(self, '_last_toggle_time'):
            if current_time - self._last_toggle_time < 200:  # 200ms防抖
                logger.info("视图切换过快，已忽略")
                return
        self._last_toggle_time = current_time

        # 切换视图模式
        if self.current_view_mode == "compact":
            self.current_view_mode = "detailed"
        else:
            self.current_view_mode = "compact"

        # 更新按钮状态
        self._update_view_toggle_button()

        # 保存视图模式到配置文件
        self._save_view_mode()

        # 调整网格间距
        if self.current_view_mode == "detailed":
            self.grid_layout.setSpacing(30)  # 详细视图使用较大间距
        else:
            self.grid_layout.setSpacing(20)  # 简约视图使用较小间距

        # 使用异步分批加载，避免卡顿
        self._load_assets_async()

        logger.info(f"切换到{'详细' if self.current_view_mode == 'detailed' else '简约'}视图")
    
    def _clear_cards(self):
        """清空所有卡片（优化版本）"""
        # 清空字典
        self.asset_cards.clear()
        
        # 批量删除，避免逐个删除的开销
        items_to_delete = []
        while self.grid_layout.count():
            child = self.grid_layout.takeAt(0)
            if child:
                widget = child.widget()
                if widget:
                    widget.setParent(None)  # 立即从父控件移除
                    items_to_delete.append(widget)

        # 延迟删除，避免阻塞UI
        from PyQt6.QtCore import QTimer
        for widget in items_to_delete:
            QTimer.singleShot(0, widget.deleteLater)
    
    def _on_scroll_position_changed(self, value: int):
        """滚动位置变化时防抖保存"""
        self._scroll_save_timer.stop()
        self._scroll_save_timer.start(500)

    def _save_scroll_position(self):
        """保存当前滚动位置到配置"""
        value = self.scroll_area.verticalScrollBar().value()
        self.controller.save_ui_state("scroll_position", value)

    def _on_assets_loaded(self, assets):
        """资产加载完成"""
        logger.info(f"Logic加载了 {len(assets)} 个资产")
        # 可以选择使用logic的资产或继续使用配置文件
    
    def _on_preview_asset(self, name):
        """预览资产"""
        logger.info(f"预览资产: {name}")

        # 查找对应的卡片和资产ID
        asset_id = None
        preview_card = None
        for aid, card in self.asset_cards.items():
            if hasattr(card, 'name') and card.name == name:
                asset_id = aid
                preview_card = card
                break

        if not asset_id or not preview_card:
            logger.error(f"未找到资产: {name}")
            return

        # 根据 package_type 分化预览行为
        pkg_type = getattr(preview_card, 'package_type', 'content')
        
        if pkg_type == 'plugin':
            # 插件不可预览
            logger.info(f"插件资产不可预览: {name}")
            return
        
        if pkg_type == 'project':
            # UE 项目：直接打开 .uproject
            self._preview_project_asset(asset_id, name)
            return
        
        if pkg_type == 'others':
            # 其他资源：检测内容类型，分别处理
            self._preview_others_asset(asset_id, name, preview_card)
            return

        # CONTENT 类型：走标准的复制预览流程（下方原有逻辑）

        # 获取可用的预览工程列表
        preview_projects = self.logic.get_additional_preview_projects_with_names()

        # 如果没有可用的预览工程，提示用户
        if not preview_projects:
            logger.warning("没有可用的预览工程，请先在设置中配置预览工程")
            # 显示提示消息框，并提供"去设置"按钮
            from .message_dialog import MessageDialog
            dialog = MessageDialog(
                "预览工程未设置",
                "请先在设置界面中添加预览工程，然后再进行资产预览。",
                "info",
                show_settings_button=True,
                parent=self
            )
            
            if dialog.exec() == MessageDialog.DialogCode.Accepted and dialog.goto_settings:
                # 用户点击了"去设置"，跳转到设置界面
                self._navigate_to_settings()
            return

        # 读取上次选择的预览工程名称
        last_selected_name = None
        try:
            user_config = self.logic.config_manager.load_user_config()
            last_selected_name = (user_config.get("last_preview_project", "")
                                    or user_config.get("last_preview_project_name", ""))
            if last_selected_name:
                logger.info(f"读取到上次选择的预览工程: {last_selected_name}")
        except Exception as e:
            logger.warning(f"读取上次选择失败: {e}")

        # 弹出预览工程选择对话框
        dialog = ProjectSelectorDialog(
            preview_projects,
            theme=self.theme,
            last_selected_name=last_selected_name,
            parent=self
        )
        if dialog.exec() != ProjectSelectorDialog.DialogCode.Accepted:
            # 用户取消了选择
            logger.info("用户取消了预览工程选择")
            return

        # 获取用户选择的工程
        selected_project = dialog.get_selected_project()
        if not selected_project:
            logger.error("未能获取选中的预览工程")
            return

        selected_name = selected_project.get("name", "")
        preview_project_path = Path(selected_project.get("path", ""))
        logger.info(f"用户选择了预览工程: {selected_name} -> {preview_project_path}")

        # 保存本次选择到配置文件
        try:
            config = self.logic.config_manager.load_user_config()
            config["last_preview_project"] = selected_name
            save_result = self.logic.config_manager.save_user_config(config, backup_reason="update_preview_project")
            if save_result:
                logger.info(f"已保存预览工程选择: {selected_name}")
            else:
                logger.warning(f"保存预览工程选择失败")
        except Exception as e:
            logger.warning(f"保存预览工程选择失败: {e}")

        # 获取预览按钮
        preview_btn = preview_card.preview_btn

        # 定义按钮重置方法
        def reset_button():
            """重置按钮状态"""
            logger.info(f"重置预览按钮: {name}")
            preview_btn.reset_progress()
            preview_btn.update_button_text("▶  预览资产")

        use_symlink_preview = False
        if self.logic and hasattr(self.logic, 'get_use_symlink_preview'):
            try:
                use_symlink_preview = self.logic.get_use_symlink_preview()
            except Exception as e:
                logger.warning(f"读取预览模式失败，默认按复制模式处理: {e}")

        if not use_symlink_preview:
            # 复制模式与导入按钮保持一致：启动时显示 0%
            preview_btn.set_progress(0.0)
            preview_btn.update_button_text("0%")

        launch_reset_scheduled = False

        def schedule_button_reset():
            """在进入启动阶段后只触发一次按钮重置。"""
            nonlocal launch_reset_scheduled
            if launch_reset_scheduled:
                return
            launch_reset_scheduled = True
            self._reset_button_signal.emit(name)

        # 根据预览模式展示不同进度：
        # - 符号链接模式：仅显示启动状态
        # - 复制模式：显示复制进度和百分比
        def update_progress(current, total, message):
            """进度回调 - 使用信号确保在主线程中更新 UI"""
            if use_symlink_preview:
                # 符号链接模式很快，只显示启动引擎相关状态
                if "启动" in message or "引擎" in message:
                    self._update_preview_text_signal.emit(name, "启动中...")
                    schedule_button_reset()
                elif current >= total and total > 0:
                    logger.info(f"资产预览准备完成: {name}")
                    schedule_button_reset()
                return

            # 复制模式：显示进度提示
            if total > 0:
                progress = max(0.0, min(1.0, current / total))
                self._update_preview_progress_signal.emit(name, progress)

            if "启动" in message or "引擎" in message:
                self._update_preview_progress_signal.emit(name, 1.0)
                self._update_preview_text_signal.emit(name, "启动中...")
                schedule_button_reset()
            elif total > 0 and current < total:
                percent = int((current / total) * 100)
                self._update_preview_text_signal.emit(name, f"{percent}%")
            elif current >= total and total > 0:
                logger.info(f"资产复制完成，准备启动引擎: {name}")
                self._update_preview_text_signal.emit(name, "启动中...")
                schedule_button_reset()

        # 直接调用logic层的预览功能（不显示初始进度）
        if self.logic:
            self.logic.preview_asset(asset_id, progress_callback=update_progress, preview_project_path=preview_project_path)

        # 监听preview_finished信号，恢复按钮状态
        def on_preview_finished():
            preview_btn.reset_progress()
            preview_btn.update_button_text("▶  预览资产")

        if self.logic:
            # 断开之前的连接（如果有）
            try:
                self.logic.preview_finished.disconnect()
            except:
                pass
            # 连接新的回调
            self.logic.preview_finished.connect(on_preview_finished)
    
    def _handle_reset_button(self, asset_name: str):
        """处理重置按钮信号（在主线程中执行）"""
        logger.info(f"主线程收到重置信号: {asset_name}")
        
        # 查找对应的卡片
        for aid, card in self.asset_cards.items():
            if hasattr(card, 'name') and card.name == asset_name:
                preview_btn = card.preview_btn
                # 在主线程中创建定时器
                QTimer.singleShot(1500, lambda: (
                    preview_btn.reset_progress(),
                    preview_btn.update_button_text("▶  预览资产")
                ))
                logger.info(f"已设置定时器重置按钮: {asset_name}")
                break
    
    def _handle_update_preview_progress(self, asset_name: str, progress: float):
        """处理更新预览进度信号（在主线程中执行）"""
        for aid, card in self.asset_cards.items():
            if hasattr(card, 'name') and card.name == asset_name:
                card.preview_btn.set_progress(progress)
                break
    
    def _handle_update_preview_text(self, asset_name: str, text: str):
        """处理更新预览文本信号（在主线程中执行）"""
        for aid, card in self.asset_cards.items():
            if hasattr(card, 'name') and card.name == asset_name:
                card.preview_btn.update_button_text(text)
                break
    
    def _preview_project_asset(self, asset_id: str, name: str):
        """预览 PROJECT 类型资产：直接打开 .uproject 文件"""
        import subprocess
        
        asset = self.logic.get_asset(asset_id)
        if not asset:
            logger.error(f"未找到资产: {asset_id}")
            return
        
        # 获取 .uproject 文件路径
        project_file_rel = getattr(asset, 'project_file', '')
        if project_file_rel:
            uproject_path = asset.path / project_file_rel
        else:
            # 尝试从 Project 子目录查找
            project_dir = asset.path / "Project"
            uproject_path = None
            if project_dir.exists():
                for f in project_dir.rglob("*.uproject"):
                    uproject_path = f
                    break
        
        if not uproject_path or not uproject_path.exists():
            from .message_dialog import MessageDialog
            MessageDialog("无法打开项目", f"未找到 .uproject 文件：{name}", "warning", parent=self).exec()
            logger.error(f"未找到 .uproject: {asset.path}")
            return
        
        logger.info(f"打开 UE 项目: {uproject_path}")
        try:
            import os
            os.startfile(str(uproject_path))
        except Exception as e:
            logger.error(f"打开项目失败: {e}", exc_info=True)
            from .message_dialog import MessageDialog
            MessageDialog("打开失败", f"无法打开项目文件：{e}", "error", parent=self).exec()
    
    def _preview_others_asset(self, asset_id: str, name: str, preview_card):
        """预览 OTHERS 类型资产：模型复制到预览工程，图片/视频/音效用 Windows 打开"""
        asset = self.logic.get_asset(asset_id)
        if not asset:
            logger.error(f"未找到资产: {asset_id}")
            return
        
        others_dir = asset.path / "Others"
        if not others_dir.exists():
            others_dir = asset.path  # 兼容旧结构
        
        # 检测内容类型
        MEDIA_EXTENSIONS = {
            '.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tga', '.tif', '.tiff',
            '.webp', '.svg', '.ico', '.exr', '.hdr',  # 图片/纹理
            '.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm',  # 视频
            '.mp3', '.wav', '.ogg', '.flac', '.aac', '.wma',  # 音频
        }
        MODEL_EXTENSIONS = {
            '.fbx', '.obj', '.gltf', '.glb', '.abc', '.usd', '.usda', '.usdc',
        }
        UE_EXTENSIONS = {'.uasset', '.umap'}
        
        has_ue_or_model = False
        has_media_only = True
        first_media_file = None
        
        try:
            for f in others_dir.rglob("*"):
                if not f.is_file():
                    continue
                ext = f.suffix.lower()
                if ext in UE_EXTENSIONS or ext in MODEL_EXTENSIONS:
                    has_ue_or_model = True
                    has_media_only = False
                    break
                if ext in MEDIA_EXTENSIONS and not first_media_file:
                    first_media_file = f
                elif ext not in MEDIA_EXTENSIONS:
                    has_media_only = False
        except Exception as e:
            logger.warning(f"扫描 Others 目录失败: {e}")
        
        if has_ue_or_model:
            # 包含模型或 UE 资产 → 走标准复制预览流程
            logger.info(f"Others 资产包含模型/UE资产，使用复制预览: {name}")
            self._on_preview_asset_content_flow(asset_id, name, preview_card)
        elif first_media_file:
            # 纯媒体文件 → 用 Windows 默认程序打开
            logger.info(f"Others 资产为媒体文件，用 Windows 打开: {first_media_file}")
            try:
                import os
                os.startfile(str(first_media_file))
            except Exception as e:
                logger.error(f"打开媒体文件失败: {e}", exc_info=True)
                from .message_dialog import MessageDialog
                MessageDialog("打开失败", f"无法打开文件：{e}", "error", parent=self).exec()
        else:
            # 无法识别内容 → 打开资产文件夹
            logger.info(f"Others 资产无法识别内容，打开文件夹: {others_dir}")
            try:
                import os
                os.startfile(str(others_dir))
            except Exception as e:
                logger.error(f"打开文件夹失败: {e}")
    
    def _on_preview_asset_content_flow(self, asset_id: str, name: str, preview_card):
        """执行 CONTENT 风格的复制预览流程（供 CONTENT 和 OTHERS 中含模型的资产使用）"""
        # 获取可用的预览工程列表
        preview_projects = self.logic.get_additional_preview_projects_with_names()

        if not preview_projects:
            logger.warning("没有可用的预览工程，请先在设置中配置预览工程")
            from .message_dialog import MessageDialog
            dialog = MessageDialog(
                "预览工程未设置",
                "请先在设置界面中添加预览工程，然后再进行资产预览。",
                "info",
                show_settings_button=True,
                parent=self
            )
            if dialog.exec() == MessageDialog.DialogCode.Accepted and dialog.goto_settings:
                self._navigate_to_settings()
            return

        # 读取上次选择的预览工程名称
        last_selected_name = None
        try:
            user_config = self.logic.config_manager.load_user_config()
            last_selected_name = (user_config.get("last_preview_project", "")
                                    or user_config.get("last_preview_project_name", ""))
        except Exception as e:
            logger.warning(f"读取上次选择失败: {e}")

        dialog = ProjectSelectorDialog(
            preview_projects,
            theme=self.theme,
            last_selected_name=last_selected_name,
            parent=self
        )
        if dialog.exec() != ProjectSelectorDialog.DialogCode.Accepted:
            return

        selected_project = dialog.get_selected_project()
        if not selected_project:
            return

        selected_name = selected_project.get("name", "")
        preview_project_path = Path(selected_project.get("path", ""))

        try:
            config = self.logic.config_manager.load_user_config()
            config["last_preview_project"] = selected_name
            self.logic.config_manager.save_user_config(config, backup_reason="update_preview_project")
        except Exception as e:
            logger.warning(f"保存预览工程选择失败: {e}")

        preview_btn = preview_card.preview_btn
        preview_btn.set_progress(0.0)
        preview_btn.update_button_text("0%")

        def reset_button():
            preview_btn.reset_progress()
            preview_btn.update_button_text("▶  预览资产")

        def update_progress(current, total, message):
            """进度回调 - 使用信号确保在主线程中更新 UI"""
            if total > 0:
                pct = current / total
                self._update_preview_progress_signal.emit(name, pct)
                self._update_preview_text_signal.emit(name, f"{int(pct * 100)}%")
            if "启动" in message or "引擎" in message:
                self._update_preview_text_signal.emit(name, "启动中...")
                self._reset_button_signal.emit(name)

        result = self.logic.preview_asset(
            asset_id,
            progress_callback=update_progress,
            preview_project_path=preview_project_path
        )
        if not result:
            reset_button()

    def _on_edit_asset_info(self, name):
        """编辑资产信息"""
        logger.info(f"编辑资产信息: {name}")
        
        try:
            # 通过控制器查找资产
            asset = self.controller.find_asset_by_name(name)
            
            if not asset:
                logger.warning(f"未找到资产: {name}")
                return
            
            # 通过控制器获取已有的资产名称和分类列表
            existing_names = self.controller.get_existing_asset_names()
            categories = self.controller.get_categories()
            
            # 检查是否有文档（文档统一存储在 .asset_config/documents/ 目录下）
            has_documentation = False
            if self.logic.documents_dir:
                doc_path = self.logic.documents_dir / f"{asset.id}.txt"
                has_documentation = doc_path.exists()
                logger.debug(f"检查文档: {doc_path}, 存在: {has_documentation}")
            
            # 导入并创建编辑对话框
            from .edit_asset_dialog import EditAssetDialog
            
            dialog = EditAssetDialog(
                logic=self.logic,
                asset_name=asset.name,
                asset_category=asset.category,
                existing_names=existing_names,
                categories=categories,
                has_documentation=has_documentation,
                parent=self
            )
            
            # 显示对话框
            if dialog.exec() == EditAssetDialog.DialogCode.Accepted:
                asset_info = dialog.get_asset_info()
                new_name = asset_info['name']
                new_category = asset_info['category']
                
                # 更新资产信息
                if self.logic.update_asset_info(
                    asset_id=asset.id,
                    new_name=new_name if new_name != asset.name else None,
                    new_category=new_category if new_category != asset.category else None
                ):
                    logger.info(f"资产 {name} 信息更新成功")
                    
                    # 只更新单个卡片，不重新加载整个列表
                    # 1. 如果分类改变，需要重新加载（因为卡片要移动到新分类）
                    if new_category != asset.category:
                        self._assets_loaded = False
                        self.load_assets_async(force_reload=True)
                        self._load_categories()
                    else:
                        # 2. 只是改名，找到对应卡片更新显示和路径即可
                        updated_asset = self.logic.get_asset(asset.id)
                        for card in self.asset_cards.values():
                            if card.name == name:
                                card.name = new_name
                                card.name_label.setText(new_name)
                                # 更新分类标签（即使分类没变，也确保显示正确）
                                if hasattr(card, 'category_label'):
                                    card.category = new_category
                                    card.category_label.setText(new_category)
                                    card.category_label.setFixedSize(card.category_label.sizeHint())
                                    # 重新定位分类标签（确保位置正确）
                                    if hasattr(card, 'thumb'):
                                        card.category_label.move(10, 153 - 10 - card.category_label.height())
                                if updated_asset and getattr(updated_asset, 'path', None):
                                    card.asset_path = str(updated_asset.path)
                                logger.info(f"已更新卡片显示与路径: {name} -> {new_name}")
                                break
                else:
                    logger.error(f"资产 {name} 信息更新失败")
                    
        except Exception as e:
            logger.error(f"编辑资产信息时出错: {e}", exc_info=True)
    
    def _on_detail_requested(self, name):
        """打开资产文档"""
        logger.info(f"[打开文档] 收到请求: {name}")
        
        try:
            # 通过控制器查找资产
            asset = self.controller.find_asset_by_name(name)
            
            if not asset:
                logger.warning(f"未找到资产: {name}")
                return
            
            logger.info(f"[打开文档] 找到资产: {asset.id}")
            
            # 检查文档目录是否存在
            if not self.logic.documents_dir:
                logger.error("文档目录未设置")
                from .message_dialog import MessageDialog
                MessageDialog("错误", "文档目录未设置", "error", parent=self).exec()
                return
            
            # 构建文档路径（.docx 格式）
            doc_path = self.logic.documents_dir / f"{asset.id}.docx"
            
            if doc_path.exists():
                # 文档存在，打开文档
                logger.info(f"[打开文档] 文档存在: {doc_path}")
                import os
                import sys
                
                if sys.platform == "win32":
                    os.startfile(str(doc_path))
                elif sys.platform == "darwin":
                    import subprocess
                    subprocess.Popen(["open", str(doc_path)])
                else:
                    import subprocess
                    subprocess.Popen(["xdg-open", str(doc_path)])
                
                logger.info(f"[打开文档] 已打开: {doc_path}")
            else:
                # 文档不存在，询问是否创建
                logger.info(f"[打开文档] 文档不存在，询问创建")
                dialog = ConfirmDialog(
                    "创建文档",
                    f"资产 \"{asset.name}\" 目前没有文档。",
                    "是否创建文档？",
                    self
                )
                if hasattr(dialog, 'center_on_parent'):
                    dialog.center_on_parent()
                
                if dialog.exec() == QDialog.DialogCode.Accepted:
                    # 创建文档
                    self._create_asset_document(asset)
            
        except Exception as e:
            logger.error(f"打开资产文档时出错: {e}", exc_info=True)
            from .message_dialog import MessageDialog
            MessageDialog("错误", f"无法打开文档：{e}", "error", parent=self).exec()
    
    def _create_asset_document(self, asset):
        """创建资产文档"""
        try:
            # 确保文档目录存在
            self.logic.documents_dir.mkdir(parents=True, exist_ok=True)
            
            # 创建文档文件（使用 .docx 格式）
            doc_path = self.logic.documents_dir / f"{asset.id}.docx"
            
            # 使用 python-docx 创建 Word 文档
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            doc = Document()

            # 标题：资产信息表（12pt）
            title = doc.add_heading('资产信息表', level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title.runs[0]
            title_run.font.size = Pt(12)

            # 添加分隔线（8pt）
            sep_para = doc.add_paragraph('=' * 50)
            sep_para.runs[0].font.size = Pt(8)

            # 基本信息区（8pt）
            doc.add_paragraph()
            info_lines = [
                f"资产名称: {asset.name}",
                f"资产ID: {asset.id}",
                f"资产类型: {asset.package_type.display_name if hasattr(asset, 'package_type') else asset.asset_type.value}",
                f"分类: {asset.category}",
                f"文件路径: {asset.path}",
                f"文件大小: {self._format_size(asset.size)}",
                f"创建时间: {asset.created_time.strftime('%Y-%m-%d %H:%M:%S') if asset.created_time else '未知'}",
            ]
            
            # 如果有引擎版本信息，添加到列表
            if hasattr(asset, 'engine_min_version') and asset.engine_min_version:
                info_lines.append(f"引擎版本: {asset.engine_min_version}")
            
            for line in info_lines:
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.size = Pt(8)

            # 分隔线（8pt）
            doc.add_paragraph()
            sep_para2 = doc.add_paragraph('=' * 50)
            sep_para2.runs[0].font.size = Pt(8)
            doc.add_paragraph()

            # 描述区（10pt 标题，8pt 内容）
            desc_heading = doc.add_heading('资产描述', level=2)
            desc_heading.runs[0].font.size = Pt(10)
            desc_content = doc.add_paragraph(asset.description or '暂无描述')
            for run in desc_content.runs:
                run.font.size = Pt(8)
            doc.add_paragraph()

            # 分隔线（8pt）
            sep_para3 = doc.add_paragraph('─' * 50)
            sep_para3.runs[0].font.size = Pt(8)
            doc.add_paragraph()

            # 使用说明区（10pt 标题，8pt 内容）
            usage_heading = doc.add_heading('使用说明', level=2)
            usage_heading.runs[0].font.size = Pt(10)
            usage_content = doc.add_paragraph('（在此添加使用说明）')
            for run in usage_content.runs:
                run.font.size = Pt(8)
            doc.add_paragraph()

            # 分隔线（8pt）
            sep_para4 = doc.add_paragraph('─' * 50)
            sep_para4.runs[0].font.size = Pt(8)
            doc.add_paragraph()

            # 注意事项区（10pt 标题，8pt 内容）
            notes_heading = doc.add_heading('注意事项', level=2)
            notes_heading.runs[0].font.size = Pt(10)
            notes_content = doc.add_paragraph('（在此添加注意事项）')
            for run in notes_content.runs:
                run.font.size = Pt(8)

            # 保存文档
            doc.save(str(doc_path))
            
            logger.info(f"[创建文档] 文档已创建: {doc_path}")
            
            # 打开文档
            import os
            import sys
            
            if sys.platform == "win32":
                os.startfile(str(doc_path))
            elif sys.platform == "darwin":
                import subprocess
                subprocess.Popen(["open", str(doc_path)])
            else:
                import subprocess
                subprocess.Popen(["xdg-open", str(doc_path)])
            
            logger.info(f"[创建文档] 已打开: {doc_path}")
            
        except Exception as e:
            logger.error(f"创建文档失败: {e}", exc_info=True)
            from .message_dialog import MessageDialog
            MessageDialog("错误", f"创建文档失败：{e}", "error", parent=self).exec()
    
    def _format_size(self, size: int) -> str:
        """格式化文件大小"""
        if size < 1024:
            return f"{size} B"
        elif size < 1024 * 1024:
            return f"{size / 1024:.1f} KB"
        elif size < 1024 * 1024 * 1024:
            return f"{size / (1024 * 1024):.1f} MB"
        else:
            return f"{size / (1024 * 1024 * 1024):.2f} GB"
    
    def _on_preview_from_dialog(self, asset_id: str):
        """从详情对话框触发的预览事件"""
        logger.info(f"[预览] 从详情对话框触发预览，资产ID: {asset_id}")
        
        try:
            asset = self.controller.find_asset_by_id(asset_id)
            
            if asset:
                logger.info(f"[预览] 找到资产名称: {asset.name}，触发预览")
                self._on_preview_asset(asset.name)
            else:
                logger.warning(f"[预览] 未找到资产ID: {asset_id}")
                
        except Exception as e:
            logger.error(f"从详情对话框预览资产时出错: {e}", exc_info=True)
    
    def _on_import_from_dialog(self, asset_id: str):
        """从详情对话框触发的导入事件"""
        logger.info(f"[导入] 从详情对话框触发导入，资产ID: {asset_id}")
        
        try:
            asset = self.controller.find_asset_by_id(asset_id)
            
            if asset:
                logger.info(f"[导入] 找到资产名称: {asset.name}，打开工程搜索窗口")
                self._on_import_asset(asset.name)
            else:
                logger.warning(f"[导入] 未找到资产ID: {asset_id}")
                
        except Exception as e:
            logger.error(f"从详情对话框导入资产时出错: {e}", exc_info=True)

    def _on_delete_asset(self, name):
        """删除资产"""
        try:
            # 通过控制器查找资产
            asset = self.controller.find_asset_by_name(name)

            if not asset:
                logger.warning(f"未找到资产: {name}")
                return

            # 显示确认对话框
            dialog = ConfirmDialog(
                "确认删除",
                f"确定要删除资产 \"{asset.name}\" 吗？",
                "注意：这将永久删除资产库中的文件/文件夹，此操作不可恢复！",
                self
            )
            
            # 居中显示
            if hasattr(dialog, 'center_on_parent'):
                dialog.center_on_parent()

            if dialog.exec() == QDialog.DialogCode.Accepted:
                # 删除资产（包括物理文件）
                success = self.logic.remove_asset(asset.id, delete_physical=True)
                if success:
                    logger.info(f"资产删除成功: {name}")
                    # 从UI中移除该资产卡片
                    if asset.id in self.asset_cards:
                        card = self.asset_cards[asset.id]
                        self.grid_layout.removeWidget(card)
                        card.deleteLater()
                        del self.asset_cards[asset.id]
                    # 重新应用筛选（保持当前分类和搜索状态）
                    self._apply_filter_to_ui()
                else:
                    logger.error(f"资产删除失败: {name}")
        except Exception as e:
            logger.error(f"删除资产时发生错误: {e}", exc_info=True)
    
    def _on_error_occurred(self, error_message: str):
        """处理错误信号，显示消息框"""
        from .message_dialog import MessageDialog
        
        # 创建消息对话框
        dialog = MessageDialog(
            "提示",
            error_message,
            "warning",
            self
        )
        dialog.exec()
        
        logger.info(f"已显示错误提示: {error_message}")
    
    def _navigate_to_settings(self):
        """跳转到设置界面"""
        try:
            # 向上查找主窗口
            main_window = None
            parent = self.parent()
            
            while parent:
                # 查找有show_settings方法的窗口（主窗口）
                if hasattr(parent, 'show_settings'):
                    main_window = parent
                    break
                parent = parent.parent()
            
            if main_window:
                logger.info("找到主窗口，准备跳转到设置界面")
                main_window.show_settings()
            else:
                logger.warning("未找到主窗口，无法跳转到设置界面")
        except Exception as e:
            logger.error(f"跳转到设置界面失败: {e}", exc_info=True)
    
    def _on_thumbnail_updated(self, asset_id: str, thumbnail_path: str):
        """处理缩略图更新信号
        
        Args:
            asset_id: 资产ID
            thumbnail_path: 新的缩略图路径
        """
        try:
            logger.info(f"收到缩略图更新信号: asset_id={asset_id}, thumbnail_path={thumbnail_path}")
            
            # 查找对应的卡片
            if asset_id in self.asset_cards:
                card = self.asset_cards[asset_id]
                
                # 从缓存中移除旧缩略图
                if asset_id in self._thumbnail_cache:
                    del self._thumbnail_cache[asset_id]
                    logger.debug(f"已清除资产 {asset_id} 的缩略图缓存")
                
                # 加载新缩略图
                if Path(thumbnail_path).exists():
                    try:
                        pixmap = QPixmap(str(thumbnail_path))
                        if not pixmap.isNull():
                            # 紧凑视图：172x115
                            scaled_pixmap = pixmap.scaled(
                                172, 115,
                                Qt.AspectRatioMode.KeepAspectRatioByExpanding,
                                Qt.TransformationMode.SmoothTransformation
                            )
                            
                            # 更新缓存
                            self._thumbnail_cache[asset_id] = scaled_pixmap
                            
                            # 更新卡片显示
                            if hasattr(card, 'thumbnail_label'):
                                card.thumbnail_label.setPixmap(scaled_pixmap)
                                logger.info(f"资产 {asset_id} 的缩略图已更新到UI")
                            else:
                                logger.warning(f"卡片 {asset_id} 没有 thumbnail_label 属性")
                    except Exception as e:
                        logger.error(f"加载新缩略图失败: {e}", exc_info=True)
                else:
                    logger.warning(f"缩略图文件不存在: {thumbnail_path}")
            else:
                logger.warning(f"未找到资产 {asset_id} 的卡片")
                
        except Exception as e:
            logger.error(f"处理缩略图更新时出错: {e}", exc_info=True)
    
    def _get_main_window(self):
        """获取顶层主窗口引用（用于驱动标题栏状态指示器）"""
        w = self.window()
        if hasattr(w, 'show_status'):
            return w
        return None
    
    def _show_add_asset_dialog(self):
        """显示添加资产流程（来源选择 -> 分析 -> 添加表单）"""
        try:
            self._start_add_asset_flow()
        except Exception as e:
            logger.error(f"显示添加资产流程时出错: {e}", exc_info=True)
    
    def _start_add_asset_flow(self, initial_path: Path = None, initial_source_type: str = None,
                              default_category: str = None):
        """统一添加资产流程：来源判定 -> 分析 -> 添加表单"""
        source_type = initial_source_type
        source_path = initial_path

        if not default_category:
            current_category = self.category_filter.currentText() if hasattr(self, 'category_filter') else "默认分类"
            if current_category in ("全部分类", "⚙️ 分类管理"):
                current_category = "默认分类"
            default_category = current_category

        if source_path is None:
            source_dialog = AddAssetSourceTypeDialog(self)
            if source_dialog.exec() != QDialog.DialogCode.Accepted:
                return

            source_type = source_dialog.selected_source
            source_path = self._pick_source_path_by_type(source_type)
            if not source_path:
                return
        else:
            if source_type is None:
                source_type = "archive" if source_path.suffix.lower() in ARCHIVE_EXTENSIONS else "folder"

        self._run_add_asset_analysis(source_path, source_type, default_category)

    def _pick_source_path_by_type(self, source_type: str):
        """根据来源类型选择源路径"""
        if source_type == "archive":
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "选择压缩包",
                "",
                "压缩包 (*.zip *.rar *.7z);;所有文件 (*)"
            )
            if not file_path:
                return None
            path = Path(file_path)
            if path.suffix.lower() not in ARCHIVE_EXTENSIONS:
                from .message_dialog import MessageDialog
                MessageDialog("不支持的文件类型", "请选择压缩包文件（.zip/.rar/.7z）", "warning", parent=self).exec()
                return None
            return path

        dir_path = QFileDialog.getExistingDirectory(
            self,
            "选择资源包文件夹",
            "",
            QFileDialog.Option.ShowDirsOnly
        )
        return Path(dir_path) if dir_path else None

    def _run_add_asset_analysis(self, source_path: Path, source_type: str, default_category: str):
        """执行识别分析详情流程"""
        password = ""
        if source_type == "archive":
            # 先尝试缓存密码
            password = ArchiveExtractor.get_cached_password() if hasattr(ArchiveExtractor, 'get_cached_password') else ""
            if ArchiveExtractor.check_password_required(source_path) and not password:
                from .password_dialog import PasswordDialog
                pwd_dialog = PasswordDialog(source_path.name, parent=self)
                if pwd_dialog.exec() != QDialog.DialogCode.Accepted:
                    return
                password = pwd_dialog.get_password()

        analysis_dialog = AddAssetAnalysisDialog(source_path.name, self)
        analysis_thread = AddAssetAnalysisThread(source_path, source_type, password=password, parent=self)
        self._current_add_analysis_thread = analysis_thread

        # 驱动标题栏进度指示器
        _mw = self._get_main_window()
        if _mw:
            _mw.show_status("资产识别分析", f"正在分析: {source_path.name}", -1)

        def _sync_stage_running(stage_idx: int, detail: str):
            analysis_dialog.set_stage_running(stage_idx, detail)
            if _mw:
                stage_names = ["分析资产版本", "分析资产类型"]
                name = stage_names[stage_idx] if stage_idx < len(stage_names) else ""
                total = len(stage_names)
                # 阶段开始：从上一阶段完成值动画到本阶段中间值（表示正在进行）
                # 阶段0: 0→25%，阶段1: 50→75%
                mid_pct = int(((stage_idx + 0.5) / total) * 100)
                if stage_idx == 0:
                    # 第一次出现，先把进度条初始化为 0
                    _mw.show_status(
                        f"阶段 {stage_idx + 1}/{total} · {name}",
                        detail,
                        0
                    )
                else:
                    _mw._status_label.setText(f"阶段 {stage_idx + 1}/{total} · {name}")
                    _mw._status_detail.setText(detail)
                _mw.animate_status_progress(mid_pct, 600)

        def _sync_stage_done(stage_idx: int):
            analysis_dialog.set_stage_done(stage_idx)
            if _mw:
                stage_names = ["分析资产版本", "分析资产类型"]
                total = len(stage_names)
                done_pct = int(((stage_idx + 1) / total) * 100)
                _mw._status_label.setText(
                    f"阶段 {stage_idx + 1}/{total} · 完成"
                )
                _mw._status_detail.setText("")
                # 平滑动画填充到该阶段完成值
                _mw.animate_status_progress(done_pct, 350)

        analysis_thread.stage_running.connect(_sync_stage_running)
        analysis_thread.stage_done.connect(_sync_stage_done)

        def _on_analysis_done(result: dict):
            analysis_dialog.result = result
            if source_type == "archive" and password:
                ArchiveExtractor.set_cached_password(password)
            if _mw:
                # 先动画到 100%，350ms 后再隐藏，避免截断动画
                _mw.animate_status_progress(100, 350)
                QTimer.singleShot(400, _mw.hide_status)
            # 展示完成态停留后自动关闭，避免一闪而过
            analysis_dialog.show_complete_then_close()

        def _on_analysis_failed(error_msg: str):
            if _mw:
                _mw.hide_status()
            analysis_dialog.reject()
            from .message_dialog import MessageDialog
            friendly = {
                "PASSWORD_REQUIRED": "压缩包需要密码，请重试并输入密码。",
                "PASSWORD_INCORRECT": "压缩包密码错误，请重试。",
                "MISSING_7Z_EXE": "7z 文件使用了不支持的压缩方法，且未找到 7-Zip 程序。",
            }
            MessageDialog("识别分析失败", friendly.get(error_msg, error_msg), "error", parent=self).exec()

        analysis_thread.analysis_done.connect(_on_analysis_done)
        analysis_thread.analysis_failed.connect(_on_analysis_failed)

        analysis_thread.start()
        accepted = analysis_dialog.exec() == QDialog.DialogCode.Accepted and analysis_dialog.result
        # 若对话框被用户手动关闭（未通过信号关闭），确保隐藏状态指示器
        if _mw:
            _mw.hide_status()
        if accepted:
            self._open_add_asset_dialog_with_analysis(
                source_path=source_path,
                source_type=source_type,
                analysis=analysis_dialog.result,
                default_category=default_category
            )

    def _open_add_asset_dialog_with_analysis(self, source_path: Path, source_type: str,
                                             analysis: dict, default_category: str):
        """分析完成后打开添加资产表单"""
        existing_names = self.controller.get_existing_asset_names()
        categories = self.controller.get_existing_categories_list()
            
        prefill_name = source_path.stem if source_type == "archive" else source_path.name
        # 压缩包优先使用原文件名，避免压缩包内部错误编码导致中文乱码
        if source_type != "archive":
            suggested_name = (analysis or {}).get("suggested_name", "")
            if suggested_name:
                prefill_name = suggested_name

        dialog = AddAssetDialog(
            existing_names,
            categories,
            parent=self,
            prefill_path=str(source_path),
            prefill_type=AssetType.PACKAGE,
            prefill_category=default_category,
            prefill_name=prefill_name,
            prefill_analysis=analysis,
            is_archive=(source_type == "archive")
        )
        dialog.center_on_parent()
            
        if dialog.exec() == AddAssetDialog.DialogCode.Accepted:
            asset_info = dialog.get_asset_info()
            logger.info(f"准备添加资产: {asset_info['name']}")
            self._add_asset_async(asset_info)
    
    def _add_asset_async(self, asset_info):
        """异步添加资产（带进度显示，支持压缩包预分析路径）"""
        from PyQt6.QtCore import QThread, pyqtSignal
        
        class AddAssetThread(QThread):
            progress_update = pyqtSignal(int, int, str)  # current, total, message
            finished = pyqtSignal(bool, object)  # success, asset
            error_message = pyqtSignal(str)  # 具体错误信息
            
            def __init__(self, logic, asset_info):
                super().__init__()
                self.logic = logic
                self.asset_info = asset_info
            
            def run(self):
                try:
                    # 确定实际添加路径
                    # 如果压缩包已预分析，使用预解压的内容路径
                    archive_content_path = self.asset_info.get('archive_content_path')
                    
                    if archive_content_path and Path(archive_content_path).exists():
                        add_path = Path(archive_content_path)
                        logger.info(f"使用预分析的压缩包内容路径: {add_path}")
                    else:
                        add_path = self.asset_info['path']
                    
                    # 导入 PackageType 用于默认值
                    from ..logic.asset_model import PackageType
                    
                    # 计算原始文件名（用于压缩包导入，避免使用临时解压目录名）
                    original_filename = ""
                    original_source = self.asset_info.get('original_source_path')
                    if original_source:
                        try:
                            original_source_path = Path(original_source)
                            original_filename = (
                                original_source_path.stem
                                if original_source_path.is_file()
                                else original_source_path.name
                            )
                        except Exception:
                            original_filename = ""
                    
                    result = self.logic.add_asset_async(
                        asset_path=add_path,
                        asset_type=self.asset_info['type'],
                        name=self.asset_info['name'],
                        category=self.asset_info['category'],
                        description="",
                        create_markdown=self.asset_info.get('create_doc', False),
                        engine_version=self.asset_info.get('engine_version', ''),
                        package_type=self.asset_info.get('package_type', PackageType.CONTENT),
                        plugin_folder_name=self.asset_info.get('plugin_folder_name', ''),
                        original_filename=original_filename,
                        progress_callback=self._progress_callback
                    )
                    
                    self.finished.emit(result is not None, result)
                except Exception as e:
                    logger.error(f"添加资产线程出错: {e}", exc_info=True)
                    self.error_message.emit(str(e))
                    self.finished.emit(False, None)
            
            def _progress_callback(self, current, total, message):
                self.progress_update.emit(current, total, message)
        
        # 保存源路径和临时目录信息，用于完成后的清理和删除确认
        self._pending_source_path = asset_info.get('original_source_path')
        self._pending_archive_extractor = asset_info.get('archive_extractor')
        self._pending_archive_temp_dir = asset_info.get('archive_temp_dir')
        self._last_add_error = ""
        
        # 创建进度对话框
        from .add_asset_progress_dialog import AddAssetProgressDialog
        progress_dialog = AddAssetProgressDialog(asset_info['name'], self)
        progress_dialog.setModal(True)
        
        # 创建并启动线程
        self.add_asset_thread = AddAssetThread(self.logic, asset_info)
        
        def _on_error_message(msg):
            self._last_add_error = msg
        
        # 连接信号
        self.add_asset_thread.progress_update.connect(progress_dialog.update_progress)
        self.add_asset_thread.error_message.connect(_on_error_message)
        self.add_asset_thread.finished.connect(self._on_add_asset_finished)
        self.add_asset_thread.finished.connect(progress_dialog.close)
        
        # 启动线程和对话框
        self.add_asset_thread.start()
        progress_dialog.show()
    
    def _on_add_asset_finished(self, success, asset):
        """添加资产完成回调"""
        # 清理压缩包临时目录
        archive_extractor = getattr(self, '_pending_archive_extractor', None)
        archive_temp_dir = getattr(self, '_pending_archive_temp_dir', None)
        if archive_extractor and archive_temp_dir:
            try:
                archive_extractor.cleanup(Path(archive_temp_dir))
            except Exception as e:
                logger.warning(f"清理临时目录失败: {e}")
        self._pending_archive_extractor = None
        self._pending_archive_temp_dir = None
        
        if success and asset:
            logger.info(f"资产添加成功: {asset.name}")
            # 刷新UI
            self._assets_loaded = False
            self.load_assets_async(force_reload=True)
            # 更新分类列表
            self._load_categories()
            
            # 根据设置自动删除源文件（不再弹窗）
            source_path = getattr(self, '_pending_source_path', None)
            if source_path and Path(source_path).exists():
                self._ask_delete_source(source_path)
        else:
            logger.error("资产添加失败")
            from .message_dialog import MessageDialog
            error_detail = getattr(self, '_last_add_error', '') or "请检查文件权限和磁盘空间。"
            MessageDialog("添加失败", f"资产添加失败：{error_detail}", "error", parent=self).exec()
        
        self._last_add_error = ""
        self._pending_source_path = None
    
    def _ask_delete_source(self, source_path):
        """根据设置决定是否删除源文件/文件夹
        
        - delete_source_after_import=True → 自动删除
        - delete_source_after_import=False（默认）→ 保留源文件，不提示
        """
        from pathlib import Path
        source_path = Path(source_path)
        if source_path.is_file():
            type_text = "压缩包"
        else:
            type_text = "文件夹"
        
        # 读取配置
        auto_delete = False
        try:
            user_config = self.logic.config_manager.load_user_config()
            auto_delete = user_config.get("delete_source_after_import", False)
        except Exception:
            pass
        
        # 仅当设置为自动删除时执行
        if auto_delete:
            try:
                if source_path.is_file():
                    source_path.unlink()
                    logger.info(f"已删除源文件: {source_path}")
                elif source_path.is_dir():
                    import shutil
                    shutil.rmtree(str(source_path))
                    logger.info(f"已删除源文件夹: {source_path}")
            except Exception as e:
                logger.error(f"删除源文件失败: {e}", exc_info=True)
                from .message_dialog import MessageDialog
                MessageDialog("删除失败", f"无法删除源文件：{e}", "error", parent=self).exec()
    
    def _show_add_category_dialog(self):
        """显示分类管理对话框"""
        try:
            from .category_management_dialog import CategoryManagementDialog
            
            # 创建对话框
            dialog = CategoryManagementDialog(self.logic, self)
            
            # 连接信号 - 分类更新后刷新列表和资产
            dialog.categories_updated.connect(self._on_categories_updated)
            
            # 显示对话框（showEvent会自动居中）
            dialog.exec()
                    
        except Exception as e:
            logger.error(f"显示分类管理对话框时出错: {e}", exc_info=True)
    
    def _on_categories_updated(self):
        """分类更新后的处理：刷新分类列表和资产显示"""
        # 刷新分类下拉列表
        self._load_categories()
        # 强制重新加载资产，确保分类标签更新
        self._assets_loaded = False
        self.load_assets_async(force_reload=True)
    
    def dragEnterEvent(self, event: QDragEnterEvent):
        """拖入事件"""
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            logger.debug("接受拖入操作")
        else:
            event.ignore()
    
    def dragMoveEvent(self, event):
        """拖动事件"""
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()
    
    def dropEvent(self, event: QDropEvent):
        """放下事件"""
        try:
            urls = event.mimeData().urls()
            if not urls:
                return
            
            file_path = Path(urls[0].toLocalFile())
            if not file_path.exists():
                logger.warning(f"拖入的路径不存在: {file_path}")
                return
            
            logger.info(f"拖入文件: {file_path}")
            
            if file_path.is_dir():
                source_type = "folder"
            elif file_path.suffix.lower() in ARCHIVE_EXTENSIONS:
                source_type = "archive"
            else:
                from .message_dialog import MessageDialog
                MessageDialog(
                    "不支持的文件类型",
                    "不支持单文件导入，请拖入文件夹或压缩包（.zip/.rar/.7z）。",
                    "info",
                    parent=self
                ).exec()
                logger.info(f"拒绝单文件拖放: {file_path}")
                return
            
            current_category = self.category_filter.currentText()
            if current_category in ("全部分类", "⚙️ 分类管理"):
                current_category = "默认分类"
            
            self._start_add_asset_flow(
                initial_path=file_path,
                initial_source_type=source_type,
                default_category=current_category,
            )
            
            event.acceptProposedAction()
            
        except Exception as e:
            logger.error(f"处理拖放操作时出错: {e}", exc_info=True)
    
    def _show_add_asset_dialog_with_prefill(self, file_path: Path, asset_type, 
                                              default_category: str, is_archive: bool = False):
        """显示添加资产对话框（预填充路径和分类）
        
        Args:
            file_path: 文件/文件夹路径
            asset_type: 资产类型
            default_category: 默认分类
            is_archive: 是否为压缩包文件
        """
        try:
            # 通过控制器获取已有的资产名称和分类列表
            existing_names = self.controller.get_existing_asset_names()
            categories = self.controller.get_existing_categories_list()
            
            # 创建对话框，传入预填充信息
            dialog = AddAssetDialog(
                existing_names, 
                categories, 
                parent=self,
                prefill_path=str(file_path),
                prefill_type=asset_type,
                prefill_category=default_category,
                is_archive=is_archive  # 直接通过构造函数传入
            )
            
            # 居中显示
            dialog.center_on_parent()
            
            # 显示对话框
            if dialog.exec() == AddAssetDialog.DialogCode.Accepted:
                asset_info = dialog.get_asset_info()
                logger.info(f"准备添加资产: {asset_info['name']}")
                
                # 异步添加资产
                self._add_asset_async(asset_info)
                    
        except Exception as e:
            logger.error(f"显示添加资产对话框时出错: {e}", exc_info=True)

    def _on_import_asset(self, name: str):
        """导入资产 - 打开工程搜索窗口"""
        logger.info(f"导入资产: {name}")
        try:
            from modules.asset_manager.ui.project_search_window import ProjectSearchWindow
            
            # 获取资产信息 - 通过遍历所有资产查找匹配的名称
            asset = None
            for a in self.logic.get_all_assets():
                if a.name == name:
                    asset = a
                    break
            
            if not asset:
                logger.error(f"未找到资产: {name}")
                return
            
            # 如果窗口已存在，先关闭
            if self.project_search_window:
                self.project_search_window.close()
                self.project_search_window = None
            
            # 创建工程搜索窗口，传递资产完整信息
            self.project_search_window = ProjectSearchWindow(
                None, 
                asset_name=name, 
                logic=self.logic,
                package_type=asset.package_type,
                engine_version=asset.engine_min_version,
                asset_path=asset.path
            )
            
            # 应用当前主题
            from core.utils.style_system import style_system
            current_theme = "modern_dark" if self.theme == "dark" else "modern_light"
            style_system.apply_to_widget(self.project_search_window, current_theme)
            
            # 设置窗口主题（更新卡片样式）
            self.project_search_window.set_theme(self.theme == "dark")
            
            # 显示窗口
            self.project_search_window.show()
            logger.info("工程搜索窗口已打开")
            
        except Exception as e:
            logger.error(f"打开工程搜索窗口时出错: {e}", exc_info=True)

    def _on_export_asset(self, name: str):
        """导出资产 - 压缩并保存到指定位置"""
        logger.info(f"导出资产: {name}")
        try:
            # 通过控制器查找资产
            asset = self.controller.find_asset_by_name(name)
            
            if not asset:
                logger.error(f"未找到资产: {name}")
                return
            
            if not asset.path or not asset.path.exists():
                logger.error(f"资产路径不存在: {asset.path}")
                from .message_dialog import MessageDialog
                MessageDialog("导出失败", "资产路径不存在", "error", parent=self).exec()
                return
            
            # 创建并显示导出进度对话框
            from modules.asset_manager.ui.export_progress_dialog import ExportProgressDialog
            
            dialog = ExportProgressDialog(name, asset.path, self)
            dialog.start_export()
            dialog.exec()
            
        except Exception as e:
            logger.error(f"导出资产时出错: {e}", exc_info=True)

    def cleanup(self):
        """清理资源"""
        try:
            # 清理缩略图缓存
            if hasattr(self, '_thumbnail_cache') and self._thumbnail_cache:
                self._thumbnail_cache.cleanup()
                logger.info("缩略图缓存已清理")
        except Exception as e:
            logger.error(f"清理资源时出错: {e}", exc_info=True)
    
    def closeEvent(self, event):
        """窗口关闭事件"""
        self.cleanup()
        super().closeEvent(event)
